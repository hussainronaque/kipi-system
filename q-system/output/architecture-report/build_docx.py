#!/usr/bin/env python3
"""Build the kipi system architecture docx report.

Reads diagram PNGs from ./diagrams/ and assembles a structured docx using
python-docx. Run from q-system/output/architecture-report/.
"""

from __future__ import annotations

import datetime as dt
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
DIAGRAMS = HERE / "diagrams"
OUTPUT = HERE / "kipi-system-architecture-2026-05-13.docx"


def set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1, color: str = "1e3a8a") -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def add_diagram(doc: Document, filename: str, caption: str, width_inches: float = 6.3) -> None:
    img_path = DIAGRAMS / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(f"Figure: {caption}")
    cap_run.font.italic = True
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_table_rows(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "1e40af")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Menlo"
    run.font.size = Pt(9)


def page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---------- TITLE PAGE ----------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Kipi System")
    tr.font.size = Pt(36)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Architecture and Implementation Report")
    sr.font.size = Pt(18)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        "An operating system built on top of Claude Code\n"
        "for running multiple businesses without losing context\n\n"
        f"Compiled {dt.date(2026, 5, 13).isoformat()}\n"
        "Author: Assaf Kipnis\n"
        "Repository: kipi-system"
    )
    mr.font.size = Pt(11)
    mr.font.italic = True

    page_break(doc)

    # ---------- TABLE OF CONTENTS ----------
    add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        ("1.", "Executive Summary"),
        ("2.", "What Problem This Solves"),
        ("3.", "Architecture Overview"),
        ("4.", "The Plugin Layer"),
        ("5.", "The Behavior Layer (Hooks, Rules, Skills, Agents)"),
        ("6.", "The Tool Layer (MCP Server and Python Harnesses)"),
        ("7.", "Rule Auto-Load Mechanism"),
        ("8.", "Memory Architecture"),
        ("9.", "Multi-Instance Architecture (KTLYST Cluster)"),
        ("10.", "PRD Operating System"),
        ("11.", "DSSE Issue Execution"),
        ("12.", "Morning Routine Pipeline"),
        ("13.", "Distribution and Propagation (kipi CLI)"),
        ("14.", "Comparison: kipi vs memory.md vs Vector Memory"),
        ("15.", "Appendix A - Complete MCP Tool Catalog"),
        ("16.", "Appendix B - File Inventory"),
    ]
    for num, label in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}\t{label}")
        run.font.size = Pt(11)

    page_break(doc)

    # ---------- 1. EXECUTIVE SUMMARY ----------
    add_heading(doc, "1. Executive Summary", level=1)
    add_para(
        doc,
        "Kipi is an operating system built on top of Claude Code. Claude Code is the runtime: "
        "context window, tool calls, MCP support, and a session loop. Kipi is the layer above that "
        "turns the runtime into a workflow system for running real businesses.",
    )
    add_para(
        doc,
        "Stated differently: Claude Code is to kipi what the Linux kernel is to a full Linux distribution. "
        "The runtime handles execution. The distribution decides how the machine behaves, what software is "
        "installed, what services run at boot, and how state is preserved across restarts.",
    )
    add_para(doc, "What kipi adds on top of Claude Code:", bold=True)

    add_table_rows(
        doc,
        ["Layer", "Count", "Purpose"],
        [
            ["Plugins", "6", "kipi-core, kipi-ops, kipi-design, kipi-dsse, prd-os, kipi-notebooklm"],
            ["MCP tools (kipi-core)", "~69", "Voice linting, lead scoring, schedule building, follow-up loops, harvest, metrics"],
            ["Python harness scripts", "30+", "Deterministic validators, schedulers, auditors, bus verifiers"],
            ["Auto-loaded rules", "19", ".claude/rules/*.md, context-aware loading"],
            ["Behavioral skills", "10+", "Voice, AUDHD, research, design, council, brand, deck-ai, NotebookLM"],
            ["Custom agents", "5", "preflight, data-ingest, content-reviewer, engagement-hitlist, synthesizer"],
            ["Hook event handlers", "6", "SessionStart, UserPromptSubmit, PreToolUse, PostToolUse, PostCompact, Stop"],
            ["Bus schemas", "35", "JSON-schema-validated payloads between pipeline agents"],
            ["Canonical knowledge files", "11", "Talk tracks, verticals, discovery, decisions, objections, market intel"],
            ["Cluster instances", "18", "Sharing state via ~/.ktlyst/bridge/"],
        ],
    )
    add_para(
        doc,
        "Kipi runs three businesses today: 4 Points Consulting, Pure Spectrum, and KTLYST Labs. "
        "Each instance has its own voice, rules, stakeholders, and deliverables. State does not leak "
        "across instances because every cross-boundary read goes through a bridge protocol.",
    )

    page_break(doc)

    # ---------- 2. WHAT PROBLEM ----------
    add_heading(doc, "2. What Problem This Solves", level=1)
    add_para(
        doc,
        "Generic LLM tools, including the Claude Code runtime in its raw form, have predictable failure modes "
        "when used as a business operations layer:",
    )

    add_table_rows(
        doc,
        ["Failure mode", "Symptom", "Kipi response"],
        [
            ["Context loss across sessions", "Decisions made on Monday are forgotten by Friday.", "Append-only canonical files, session-start memory load, handoff doc, auto-memory index."],
            ["Voice drift", "AI produces text that does not sound like the founder; AI-detector patterns leak through.", "founder-voice skill, voice-enforcement rule, kipi_voice_lint deterministic checker, anti-AI banned-word lists."],
            ["Role bleed across businesses", "Strategy advice contaminated by product details from a different company.", "Per-instance scoped rules, multi-instance state bridges with explicit reader/writer contracts."],
            ["LLM hallucination on numeric work", "Lead scores, A/B math, churn signals all made up.", "Deterministic Python scorers: kipi_score_lead, kipi_ab_test_calc, kipi_churn_health_score. No LLM in the validation loop."],
            ["Lost follow-ups", "Conversations end with open loops that never close.", "loop_open/loop_close/loop_escalate MCP tools, daily loop review agent."],
            ["Sycophancy and unflagged opinion drift", "AI agrees with the founder and never challenges.", "council skill (multi-persona debate), sycophancy rule, decision-origin tagging."],
            ["No discipline for product work", "Code changes ship without scope, review, or verification receipts.", "prd-os + kipi-dsse: gated PRD workflow, scope enforcement on edits, Codex adversarial review, signed receipts."],
            ["Executive function overload", "Walls of text, demand language, lost actions for an AUDHD founder.", "AUDHD output style, audhd-interaction rule, executive-function skill applied to every actionable output."],
        ],
    )

    page_break(doc)

    # ---------- 3. ARCHITECTURE OVERVIEW ----------
    add_heading(doc, "3. Architecture Overview", level=1)
    add_para(
        doc,
        "Kipi is structured in four layers stacked on top of the Claude Code runtime. Each layer is independently "
        "replaceable. The plugin layer composes capabilities. The behavior layer shapes how the model responds. "
        "The tool layer adds deterministic logic. The data layer is the durable substrate that survives sessions, "
        "compactions, and machine restarts.",
    )
    add_diagram(doc, "01-architecture-layers.png", "Kipi system architecture, four layers above Claude Code")

    add_heading(doc, "3.1 The Layer Contract", level=2)
    add_para(
        doc,
        "The plugin layer is the unit of distribution. Each plugin has its own plugin.json, optional MCP server, "
        "and bundles of skills, commands, and hooks. Plugins are versioned independently (semver) and can be "
        "enabled or disabled per instance. The plugins map to user-facing capability groups, not to internal "
        "code organization.",
    )
    add_para(
        doc,
        "The behavior layer is loaded into the model context. Rules live in .claude/rules and are auto-attached "
        "by Claude Code based on path globs or always-on flags. Skills live in plugins and are invoked explicitly "
        "or via auto-detection rules. Agents are spawned as sub-conversations with their own model selection, "
        "tool allowlist, and effort setting.",
    )
    add_para(
        doc,
        "The tool layer is where determinism lives. The kipi-mcp server registers around 69 MCP tools; the "
        "majority of these wrap pure Python harness functions that compute pass-or-fail results without an "
        "LLM in the loop. This is the layer that makes the difference between an opinionated LLM assistant "
        "and a system the founder can trust for numeric or rule-bound decisions.",
    )
    add_para(
        doc,
        "The data layer is multi-store. Markdown files hold human-readable canonical knowledge. JSON files hold "
        "structured working memory and bus payloads. A SQLite database (metrics_store) holds append-only event "
        "data with deterministic query shapes. Cross-instance state lives in ~/.ktlyst/bridge/ where each file "
        "has exactly one writer and a known set of readers.",
    )

    page_break(doc)

    # ---------- 4. PLUGIN LAYER ----------
    add_heading(doc, "4. The Plugin Layer", level=1)
    add_para(
        doc,
        "Kipi composes six plugins. Each is shipped as a directory under plugins/ with a .claude-plugin/plugin.json "
        "manifest. The plugins together cover voice and content, design and brand, GTM and council, gated product "
        "workflows, and external integrations.",
    )
    add_diagram(doc, "02-plugin-map.png", "Plugin landscape, skills (yellow), commands (green), hooks (pink)")

    add_table_rows(
        doc,
        ["Plugin", "Version", "Surface area"],
        [
            ["kipi-core", "1.3.0", "Founder voice, AUDHD executive function, research mode, LinkedIn brand, deck-ai, plus kipi-mcp server with ~69 tools."],
            ["kipi-ops", "1.2.0", "council skill (multi-persona debate with quick + deliberation workflows)."],
            ["kipi-design", "1.2.0", "brand, design (CIP/logo/slides/icons/social), ui-ux-pro-max (shadcn + tailwind + 99 UX guidelines)."],
            ["kipi-dsse", "0.1.0", "DSSE issue execution: scope_hook + stop_gate, six issue-* commands, JSONL findings."],
            ["prd-os", "0.1.0", "PRD lifecycle: six prd-* commands, scope_hook + stop_gate, Codex review integration."],
            ["kipi-notebooklm", "0.1.0", "NotebookLM MCP wrapper. Drives real Chrome via Patchright. Personal Gmail profile."],
        ],
    )

    add_heading(doc, "4.1 Why Six Plugins, Not One Monolith", level=2)
    add_para(
        doc,
        "Different instances need different plugin sets. A consulting instance does not need prd-os. A legal "
        "instance does not need kipi-design. By splitting along capability lines, the founder enables exactly "
        "the plugins each role-specific instance needs, without forcing every instance to carry every dependency.",
    )

    page_break(doc)

    # ---------- 5. BEHAVIOR LAYER ----------
    add_heading(doc, "5. The Behavior Layer", level=1)
    add_para(
        doc,
        "The behavior layer is everything that shapes how the model responds before, during, and after a tool call. "
        "It is implemented through four mechanisms: hooks (deterministic shell or Python scripts that fire on "
        "lifecycle events), rules (markdown files that auto-load into context), skills (markdown bundles invoked "
        "explicitly), and agents (sub-conversations with their own model, tools, and prompt).",
    )

    # 5.1 hooks
    add_heading(doc, "5.1 Hooks", level=2)
    add_para(
        doc,
        "Hooks are wired in .claude/settings.json. They run as subprocess calls at well-defined points in the "
        "Claude Code session lifecycle. Their job is to mutate context, enforce gates, or log state outside the "
        "model's reasoning.",
    )
    add_diagram(doc, "03-hook-lifecycle.png", "Session hook lifecycle, from SessionStart through Stop")

    add_table_rows(
        doc,
        ["Event", "Script", "What it does"],
        [
            ["SessionStart", "git-health-check.sh", "Prints the fleet git status banner (dirty repos, ahead/behind, branches)."],
            ["SessionStart", "session-start.py", "Loads memory index, resolves paths, surfaces any missed wraps."],
            ["SessionStart", "md-prune.py", "Archives stale markdown beyond freshness thresholds."],
            ["UserPromptSubmit", "token-guard.py", "Hard stop guard for over-long autonomous runs."],
            ["PreToolUse (any)", "token-guard.py", "Same guard, evaluated before each tool call."],
            ["PostToolUse (Edit, Write)", "wiring-check.py", "Verifies every file change is connected end-to-end across plugins, hooks, MCP, agents, bus, canonical, rules."],
            ["PostCompact", "post-compact.sh", "Re-injects mode, open loops, and voice reminders after context compaction."],
            ["Stop", "auto-commit.py (async)", "Commits dirty files at session end."],
            ["Stop", "stop-logger.sh (async)", "Appends session metadata to a structured log."],
        ],
    )

    # 5.2 rules
    add_heading(doc, "5.2 Rules", level=2)
    add_para(
        doc,
        "Rules live in .claude/rules/. Each rule is a markdown file with optional frontmatter declaring auto-load "
        "conditions (paths, globs, always-on). Claude Code injects matching rules into the model context before "
        "each response. Rules do not run; they instruct.",
    )
    add_table_rows(
        doc,
        ["Rule file", "Purpose"],
        [
            ["anti-misclassification.md", "Guardrails preventing content from being misclassified or misrouted."],
            ["audhd-interaction.md", "AUDHD and ADHD-aware interaction patterns for all founder-facing output."],
            ["auto-detection.md", "Triggers for pasted transcripts, social posts, and council auto-invoke without explicit commands."],
            ["coding-audhd.md", "AUDHD-adapted coding rules: structure, communication, emotional scaffolding."],
            ["coding-standards.md", "Naming and style conventions for Python, JS, Shell, and JSON."],
            ["content-output.md", "Output format rules for generated marketing content."],
            ["design-auto-invoke.md", "Loads kipi-design skills only for public-facing assets."],
            ["dev-skills-auto-invoke.md", "Loads dev skills for plugin, MCP, hook, and Claude API work."],
            ["folder-structure.md", "Strict placement enforcement for kipi-system."],
            ["marketing-system.md", "Pipeline rules for marketing content lifecycle."],
            ["md-hygiene.md", "Line budgets, section pinning, auto-pruning for canonical files."],
            ["memory-freshness.md", "Decay rules: fast/medium/slow files and their thresholds."],
            ["morning-pipeline.md", "Execution rules for the nine-phase morning routine."],
            ["security.md", "Secret handling, dangerous-operation deny list."],
            ["social-reaction-gate.md", "ENFORCED: extract poster's claims before drafting any comment, reply, or DM."],
            ["sycophancy.md", "Decision-origin tagging and sycophancy awareness."],
            ["token-discipline.md", "Token consumption guardrails and self-monitoring."],
            ["voice-enforcement.md", "Founder voice applied to all external-facing written content."],
            ["wiring-check.md", "ENFORCED: definition of done. No task is done until every change is wired end-to-end."],
        ],
    )

    page_break(doc)

    # 5.3 skills
    add_heading(doc, "5.3 Skills", level=2)
    add_para(
        doc,
        "Skills are markdown bundles, each with a SKILL.md entrypoint and optional references/ and templates/ "
        "subdirectories. Unlike rules (which auto-attach), skills are invoked explicitly through the Skill tool "
        "or via auto-detection rules. They are designed for heavyweight transformations that justify the cost "
        "of loading their full reference content.",
    )
    add_table_rows(
        doc,
        ["Plugin", "Skill", "What it does"],
        [
            ["kipi-core", "founder-voice", "Voice profile + 20+ writing samples + 5 archetypes. Loaded for every external-facing piece of writing."],
            ["kipi-core", "audhd-executive-function", "Output formatting for actionable founder output: max-3 bullets, copy-paste-ready, energy-tagged."],
            ["kipi-core", "research-mode", "Anti-hallucination research mode with citation grounding and 'I don't know' behavior."],
            ["kipi-core", "linkedin-brand", "LinkedIn playbook, voice-check, summary frameworks, LLM visibility tracking."],
            ["kipi-core", "deck-ai", "Slidev-based modern deck generation, layout catalog, Unsplash imagery."],
            ["kipi-ops", "council", "Multi-persona debate (quick or full). Auto-fires on significant canonical changes."],
            ["kipi-design", "brand", "Visual identity, voice framework, color palette management, typography, logo usage."],
            ["kipi-design", "design", "Logo + CIP + banners + slides + icons + social photos + design tokens."],
            ["kipi-design", "ui-ux-pro-max", "shadcn + tailwind + 99 UX guidelines + 161 palettes + 57 font pairings + 25 chart types."],
        ],
    )

    # 5.4 agents
    add_heading(doc, "5.4 Agents", level=2)
    add_para(
        doc,
        "Agents are sub-conversations spawned by the orchestrator. Each has its own model, tool allowlist, and "
        "system prompt. Agent definitions live in .claude/agents/*.md as YAML frontmatter plus a markdown body.",
    )
    add_table_rows(
        doc,
        ["Agent", "Model", "Effort", "Purpose"],
        [
            ["preflight", "claude-haiku-4-5", "default", "Verify tools and files available before morning routine; writes preflight.json."],
            ["data-ingest", "claude-haiku-4-5", "default", "Pull calendar, email, Notion data. Structured extraction only, no analysis."],
            ["content-reviewer", "claude-sonnet-4-6", "default", "Four-pass review on any content before send: voice, guardrails, anti-AI, actionability."],
            ["engagement-hitlist", "claude-opus-4-6", "max", "Generate ranked, copy-paste-ready engagement actions from pipeline data."],
            ["synthesizer", "claude-opus-4-6", "max", "Final assembly: reads all bus data, writes daily schedule HTML."],
        ],
    )
    add_para(
        doc,
        "Model selection is meaningful. Cheap haiku models handle deterministic extraction. Sonnet handles "
        "review with judgment. Opus handles synthesis and ranking where reasoning depth materially changes "
        "the output quality.",
    )

    page_break(doc)

    # ---------- 6. TOOL LAYER ----------
    add_heading(doc, "6. The Tool Layer", level=1)
    add_para(
        doc,
        "The tool layer is where most of the deterministic logic lives. There are two parts: the kipi-mcp MCP "
        "server that registers around 69 tools across categorized domains, and an additional 30+ standalone "
        "Python scripts that run as hook subprocesses or are invoked directly by harnesses.",
    )

    add_heading(doc, "6.1 The kipi-mcp Server", level=2)
    add_para(
        doc,
        "The kipi-mcp server (plugins/kipi-core/kipi-mcp/src/kipi_mcp/server.py) is a single Python module of "
        "around 1,694 lines that registers tools via @mcp.tool() decorators. The server imports specialized "
        "modules for the heavy lifting: linter.py, scorer.py, schedule_verifier.py, validator.py, harvest_store.py, "
        "loop_tracker.py, metrics_store.py, schema_gen.py, morning_init.py, morning_auditor.py, draft_scanner.py, "
        "orchestrator_verifier.py, bus_verifier.py, harvest_orchestrator.py, step_logger.py, and others.",
    )
    add_para(
        doc,
        "Each MCP tool is a thin shim that validates inputs, calls a harness function, and returns a JSON-serialized "
        "result. The harness functions are pure Python with no LLM involvement. This is the single most important "
        "architectural fact about kipi: numeric and rule-bound work does not flow through model inference. It "
        "flows through code.",
    )
    add_diagram(doc, "04-mcp-tool-flow.png", "MCP tool call lifecycle, model never participates in the verification step")

    add_heading(doc, "6.2 Tool Categories", level=2)
    add_para(doc, "The ~69 kipi-core tools cluster into ten functional groups:")
    add_table_rows(
        doc,
        ["Group", "Tools", "Representative responsibilities"],
        [
            ["Instance management", "kipi_suggest_instance_name, kipi_set_instance_name, kipi_session_bootstrap, kipi_session_handoff, kipi_preflight, kipi_canonical_digest, kipi_morning_init", "Naming, bootstrap, session continuity, paths, digest generation."],
            ["Validation", "kipi_validate, kipi_validate_schedule, kipi_validate_ad_copy, kipi_validate_cold_email, kipi_seo_check, kipi_voice_lint, kipi_copy_edit_lint, kipi_linkedin_gate", "Deterministic pass/fail for content before it ships."],
            ["Scoring + math", "kipi_score_lead, kipi_ab_test_calc, kipi_churn_health_score, kipi_cancel_flow_offer, kipi_crack_detect", "Pure Python math. No LLM in the loop."],
            ["Morning step logging", "log_init, log_step, log_add_card, log_deliver_cards, log_gate_check, log_checksum, log_verify", "Append-only step ledger for the morning routine."],
            ["Follow-up loops", "loop_open, loop_close, loop_force_close, loop_escalate, loop_touch, loop_prune", "Cross-session loop registry with escalation logic."],
            ["Content building", "kipi_create_template, kipi_build_schedule, kipi_scan_draft, kipi_generate_schema", "Schedule HTML, template synthesis, JSON-LD generation."],
            ["Bus and orchestration", "kipi_verify_schedule, kipi_verify_bus, kipi_verify_orchestrator, kipi_bus_to_log, kipi_audit_morning, kipi_gate_check, kipi_deliverables_check", "Pipeline integrity verifiers."],
            ["Metrics database", "kipi_init_db, kipi_insert_content_metrics, kipi_insert_behavioral_signals, kipi_insert_outreach, kipi_insert_copy_edit, kipi_query, kipi_daily_metrics, kipi_monthly_learnings, kipi_log_agent_metric, kipi_agent_metrics", "SQLite append-only metrics store with query shapes."],
            ["Harvest", "kipi_store_harvest, kipi_get_harvest, kipi_harvest_status, kipi_harvest_summary, kipi_harvest_cleanup, kipi_harvest_health, kipi_approve_apify_budget", "External data harvest (Apify, etc.) with budget gating."],
            ["LinkedIn cadence", "kipi_log_linkedin_activity, kipi_linkedin_cadence_check, kipi_linkedin_gate", "LinkedIn rhythm enforcement and gating."],
            ["Notion queue", "kipi_queue_notion_write, kipi_get_notion_queue", "Decoupled Notion write queue with source-agent attribution."],
            ["Backup and portability", "kipi_backup, kipi_export, kipi_import", "Snapshot, export, import for data portability."],
        ],
    )

    page_break(doc)

    add_heading(doc, "6.3 Resources", level=2)
    add_para(
        doc,
        "In addition to tools, the kipi-mcp server exposes five MCP resources:",
    )
    add_code(
        doc,
        "kipi://paths        - resolved directory paths\n"
        "kipi://status       - server status and version\n"
        "kipi://instances    - registered cluster instances\n"
        "kipi://loops/open   - currently open follow-up loops\n"
        "kipi://loops/stats  - aggregate loop statistics\n"
        "kipi://backups      - available backup archives",
    )

    add_heading(doc, "6.4 Why Determinism Matters", level=2)
    add_para(
        doc,
        "Three places in the system would break catastrophically if they were LLM-generated:",
    )
    add_para(
        doc,
        "1. Lead scoring. A score of 78 must mean the same thing today and a month from now. An LLM's "
        "stochastic scoring drifts; the deterministic scorer in scorer.py loads weights from a registry and "
        "applies them with arithmetic.",
    )
    add_para(
        doc,
        "2. Voice linting. The banned-words and anti-AI pattern lists are hard rules. A model asked to "
        "self-check its own output for AI tells has an obvious conflict of interest. The lint runs as Python "
        "regex against the rendered text.",
    )
    add_para(
        doc,
        "3. Bus verification. The morning pipeline produces JSON payloads against 35 schemas. Whether a "
        "payload validates is a binary fact, not a judgment call. The bus_verifier walks every schema and "
        "every payload, returning structured violations.",
    )

    page_break(doc)

    # ---------- 7. RULE AUTOLOAD ----------
    add_heading(doc, "7. Rule Auto-Load Mechanism", level=1)
    add_para(
        doc,
        "Rules attach to context conditionally. The decision happens before each model response. A rule can "
        "match on file path, content pattern, or always-on flag. Some rules fire only after the model has "
        "decided to use a specific tool (the wiring-check rule fires only on Edit or Write).",
    )
    add_diagram(doc, "05-rule-autoload.png", "Rule auto-load decision flow, from prompt to response")

    add_para(
        doc,
        "Auto-detection rules go further. They watch for content patterns in the user's prompt and trigger "
        "downstream workflows without an explicit command. The most important examples:",
    )
    add_table_rows(
        doc,
        ["Trigger pattern", "Rule", "Workflow"],
        [
            ["Pasted conversation transcript or meeting notes", "auto-detection.md", "Auto-fires /q-debrief: extracts the person, role, company; runs all 12 strategic implications lenses; routes outputs to canonical files; logs to Notion."],
            ["Social post screenshot", "auto-detection.md", "Routes to /q-engage reactive mode: market-intel evaluation + best-comment generation."],
            ["Significant canonical change (>5 lines or new section)", "auto-detection.md", "Auto-fires council quick mode; if 2+ personas object, dissent is surfaced before writing."],
            ["Comment, reply, DM on someone else's content", "social-reaction-gate.md", "ENFORCED extraction of poster's claims; founder confirms before any draft."],
            ["External-facing content (post, email, DM)", "voice-enforcement.md", "Loads founder-voice skill, applies anti-AI patterns, banned-word lists."],
            ["Edit or Write touching plugins, MCP tools, agents", "wiring-check.md", "Verifies every change is connected end-to-end before declaring done."],
        ],
    )

    page_break(doc)

    # ---------- 8. MEMORY ----------
    add_heading(doc, "8. Memory Architecture", level=1)
    add_para(
        doc,
        "Kipi memory is not a single file. It is a layered substrate, each layer with its own lifecycle, "
        "freshness rules, and access patterns. The layers are loaded conditionally based on what the founder "
        "is doing.",
    )
    add_diagram(doc, "06-memory-layers.png", "Memory layers, from global preferences to ephemeral working state")

    add_heading(doc, "8.1 Global Memory", level=2)
    add_para(
        doc,
        "Loaded every session for every project. Contains the user's role, preferences, AUDHD profile, and the "
        "auto-memory index. Auto-memory itself is a curated set of named markdown files in "
        "~/.claude/projects/*/memory/, each with frontmatter that declares its type (user, feedback, project, "
        "reference) and links to related memories via [[name]] syntax.",
    )

    add_heading(doc, "8.2 Project Memory", level=2)
    add_para(
        doc,
        "Loaded per project. Includes the project CLAUDE.md, q-system/CLAUDE.md, and all matching .claude/rules. "
        "Rules attach conditionally. Some rules are always-on (audhd-interaction, voice-enforcement, "
        "wiring-check). Others fire only on specific globs (folder-structure on file touches inside the "
        "project, design-auto-invoke on public-facing assets).",
    )

    add_heading(doc, "8.3 Canonical Layer", level=2)
    add_para(
        doc,
        "Eleven structured knowledge files in q-system/canonical/. These are the durable answers to the "
        "questions the founder will be asked repeatedly: 'how do you position?' goes to talk-tracks.md. "
        "'what verticals do you go after?' goes to verticals.md. 'what do you charge?' goes to pricing-framework.md. "
        "These files are written once per insight and updated explicitly via /q-calibrate, never by accident.",
    )
    add_para(
        doc,
        "Crucially, edits to canonical files trigger the wiring-check rule and, for significant changes, "
        "auto-invoke the council skill for dissent detection. Canonical is durable, but it is not blindly "
        "trusted - the system has guardrails against the founder being talked into bad changes.",
    )

    add_heading(doc, "8.4 Working Memory", level=2)
    add_para(
        doc,
        "Ephemeral per-session or per-day state. The morning log (morning-log-YYYY-MM-DD.json) tracks every "
        "step of today's pipeline. The bus directory (agent-pipeline/bus/) holds today's pipeline payloads. "
        "memory/last-handoff.md is the bridge between sessions. The loops registry tracks open follow-ups.",
    )

    add_heading(doc, "8.5 Append-Only Database", level=2)
    add_para(
        doc,
        "metrics_store.py implements a SQLite database with four tables: content_metrics, behavioral_signals, "
        "outreach, copy_edits. Every insert is append-only. Queries are predefined (kipi_query, kipi_daily_metrics, "
        "kipi_monthly_learnings) - the database is not a free-form analytics surface, it is a structured event "
        "log with known reporting shapes.",
    )

    add_heading(doc, "8.6 Cross-Instance Memory: The Bridge", level=2)
    add_para(
        doc,
        "Shared state across the KTLYST cluster lives in ~/.ktlyst/bridge/. Every bridge file has exactly one "
        "writer and a defined set of readers. The protocol is documented in q-system/CLAUDE.md and enforced by "
        "the ktlyst-cluster.md rule that loads in every cluster instance.",
    )
    add_table_rows(
        doc,
        ["Bridge file", "Writer", "Readers", "Contents"],
        [
            ["product_state.json", "product", "strategy", "Pipeline capabilities, demo status, release-readiness verdict."],
            ["canonical-digest.json", "strategy", "product, website, lawyer", "Talk tracks summary, positioning, anti-misclassification."],
            ["legal-flags.json", "lawyer", "strategy", "Compliance flags, contract status, legal blockers."],
            ["website-state.json", "website", "strategy", "Deployed copy version, A/B test results."],
            ["market_signal.json", "strategy", "all", "Market signals, competitive intel."],
            ["threat_status_history.json", "product", "strategy", "Threat processing pipeline status."],
        ],
    )

    page_break(doc)

    # ---------- 9. MULTI INSTANCE ----------
    add_heading(doc, "9. Multi-Instance Architecture (KTLYST Cluster)", level=1)
    add_para(
        doc,
        "Kipi is built to run multiple businesses simultaneously without state contamination. The KTLYST cluster "
        "is the canonical example, but the architecture generalizes. The skeleton (kipi-system) is the source of "
        "truth for shared infrastructure. Each instance is a separate git repo that the kipi CLI keeps in sync "
        "with the skeleton.",
    )
    add_diagram(doc, "07-multi-instance-bridge.png", "Multi-instance topology with shared bridge state")

    add_heading(doc, "9.1 Instance Roles in the KTLYST Cluster", level=2)
    add_table_rows(
        doc,
        ["Instance", "Role", "Canonical authority"],
        [
            ["ktlyst-strategy", "GTM, fundraising, positioning, relationships", "Positioning, talk tracks, objections, market intelligence."],
            ["ktlyst-product", "Python extraction pipeline, demos, deliverables", "Product capabilities, technical architecture."],
            ["ktlyst-website", "Marketing site (React + Vite + Tailwind, Vercel)", "Deployed marketing copy."],
            ["ktlyst-lawyer", "Legal counsel, compliance, contracts", "Legal frameworks, compliance, IP, contract templates."],
            ["ktlyst-personal-brand", "Founder LinkedIn brand and content", "Personal brand voice and visibility."],
        ],
    )

    add_heading(doc, "9.2 Authority Rules", level=2)
    add_para(
        doc,
        "Strategy owns positioning. Product owns technical truth. Lawyer owns legal frameworks. Website reflects, "
        "never defines. These authority rules are enforced by the ktlyst-cluster.md rule that loads in every "
        "instance and prevents cross-domain contradictions.",
    )

    add_heading(doc, "9.3 Other Instances on the Same Architecture", level=2)
    add_para(
        doc,
        "The same skeleton runs multiple businesses outside the KTLYST cluster: 4_points_consulting, "
        "Pure_spectrum_Q, ASK_AI_consultant, and others. Each gets its own rules, skills, voice, stakeholders, "
        "and deliverables. The skeleton propagates infrastructure; instance-specific content stays in the instance.",
    )

    page_break(doc)

    # ---------- 10. PRD OS ----------
    add_heading(doc, "10. PRD Operating System", level=1)
    add_para(
        doc,
        "prd-os is a formal, gated workflow for capturing a rough idea, drafting a PRD, running adversarial AI "
        "review, triaging findings, and decomposing the approved PRD into atomic issue specs. It is the system "
        "used for any product or system change in kipi.",
    )
    add_diagram(doc, "08-prd-os-state.png", "PRD state machine: idea -> draft -> in_review -> approved -> archived")

    add_heading(doc, "10.1 The Six PRD Commands", level=2)
    add_table_rows(
        doc,
        ["Command", "Effect"],
        [
            ["/prd-start <idea>", "Captures a rough idea, scaffolds the PRD spec, blocks if an issue is already in-progress."],
            ["/prd-review", "Runs Codex review + adversarial review. Streams findings to JSONL."],
            ["/prd-triage", "Founder triages every pending finding with a disposition."],
            ["/prd-approve", "Advances PRD to approved. Blocked by any pending finding."],
            ["/prd-split", "Decomposes the approved PRD into one issue spec per manifest entry."],
            ["/prd-archive", "Final archive. Blocked until every accepted finding has a receipt."],
        ],
    )

    add_heading(doc, "10.2 Non-Negotiables", level=2)
    add_para(
        doc,
        "PRD drafting cannot drift into implementation. Scope enforcement restricts edits to the PRD file "
        "during drafting. Concurrent PRD and issue contexts are blocked. Codex never edits - it returns findings "
        "for Claude to triage. Claude is the sole editor. Every finding gets one of four dispositions: "
        "must-fix, optional, deferred, or rejected-with-reason. No finding may stay unset.",
    )

    page_break(doc)

    # ---------- 11. DSSE ----------
    add_heading(doc, "11. DSSE Issue Execution", level=1)
    add_para(
        doc,
        "kipi-dsse executes the issue specs produced by prd-os. Its job is scope enforcement, verification, "
        "and signed receipts. It is the layer that prevents an issue from drifting into a refactor.",
    )
    add_diagram(doc, "09-issue-dsse-state.png", "Issue state machine: open -> in_progress -> verified -> reviewed -> closed")

    add_heading(doc, "11.1 The Scope Hook", level=2)
    add_para(
        doc,
        "plugins/kipi-dsse/hooks/scope_hook.py runs on every Edit or Write. It reads the active issue's "
        "allowed_files list from .claude/state/active-issue.json and rejects any edit outside that list. "
        "Empty allowed_files means deny-all except the spec itself.",
    )

    add_heading(doc, "11.2 The Stop Gate", level=2)
    add_para(
        doc,
        "plugins/kipi-dsse/hooks/stop_gate.py fires on session stop. It checks for an in-progress issue and "
        "either honors a stop_hook_active flag from Claude Code, or falls back to signature-based exhaustion "
        "(max 3 firings per signature, tracked in .claude/state/stop-gate-firings.json).",
    )

    add_heading(doc, "11.3 Receipts", level=2)
    add_para(
        doc,
        "Three receipts must exist between approve and close:",
    )
    add_table_rows(
        doc,
        ["Receipt", "Written by", "Means"],
        [
            ["verified", "/issue-verify", "All required_checks for the issue have passed."],
            ["reviewed", "/issue-review", "Codex native and adversarial review have run, scoped to allowed_files."],
            ["findings_triaged", "/issue-closeout", "Every Codex finding has a disposition."],
        ],
    )

    page_break(doc)

    # ---------- 12. MORNING ----------
    add_heading(doc, "12. Morning Routine Pipeline", level=1)
    add_para(
        doc,
        "The morning routine is kipi's largest single workflow. It runs nine phases of agents, writes ~30 bus "
        "payloads to disk, and produces the founder's daily schedule HTML. It is the workflow that proves the "
        "architecture: every plugin, every layer, every memory store participates.",
    )
    add_diagram(doc, "10-morning-pipeline.png", "Morning routine, nine phases, agent-pipeline bus topology")

    add_heading(doc, "12.1 Phase Breakdown", level=2)
    add_table_rows(
        doc,
        ["Phase", "Purpose", "Bus output"],
        [
            ["00 - Preflight", "Verify tools and files; gate the pipeline.", "preflight.json"],
            ["01 - Ingest", "Pull calendar, gmail, CRM, Notion, copy-diff.", "calendar.json, gmail.json, crm.json, etc."],
            ["02 - Enrich", "Meeting prep, warm-intro match, X activity.", "meeting-prep.json, warm-intros.json, x-activity.json"],
            ["03 - Content intel", "Content intelligence, LinkedIn DMs.", "content-intel.json, linkedin-dms.json"],
            ["04 - Brand", "Founder brand post, post visuals.", "founder-brand-post.json, post-visuals.json"],
            ["05 - Pipeline", "Pipeline follow-up, engagement hitlist, lead sourcing, loop review.", "pipeline-followup.json, hitlist.json, leads.json"],
            ["06 - Deliverables", "Client deliverables status.", "client-deliverables.json"],
            ["07 - Synthesize", "Build the daily schedule HTML from all bus data.", "schedule-data.json + HTML"],
            ["08 - Verify", "Visual verification, layout sanity.", "publish-reconciliation.json"],
            ["09 - Push", "Push to Notion and CRM.", "notion.json, crm.json"],
        ],
    )
    add_para(
        doc,
        "Every bus payload is validated against its JSON schema before any downstream agent reads it. If "
        "validation fails, the pipeline halts at that phase and surfaces the violation to the founder.",
    )

    page_break(doc)

    # ---------- 13. DISTRIBUTION ----------
    add_heading(doc, "13. Distribution and Propagation (kipi CLI)", level=1)
    add_para(
        doc,
        "The kipi CLI is a single bash script (./kipi) that wraps the propagation logic. It is installed as "
        "/opt/homebrew/bin/kipi and used from any directory. The skeleton lives at the kipi-system repo; "
        "instances are the projects registered in instance-registry.json.",
    )
    add_diagram(doc, "11-kipi-update-propagation.png", "kipi update propagation across registered instances")

    add_heading(doc, "13.1 CLI Surface", level=2)
    add_table_rows(
        doc,
        ["Command", "What it does"],
        [
            ["kipi update", "Pull latest skeleton into all registered instances."],
            ["kipi update --dry", "Preview without changing anything."],
            ["kipi new <path> <name>", "Create a new project with the skeleton baked in."],
            ["kipi push", "Push generic improvements from current project back to skeleton."],
            ["kipi dev", "Launch Claude Code with all kipi plugins loaded."],
            ["kipi check [phase]", "Run validation harness (validate-separation.py)."],
            ["kipi migrate <path>", "Migrate an instance to full kipi compliance."],
            ["kipi cluster add <path> <name> <role>", "Plug an instance into the KTLYST cluster."],
            ["kipi cluster sync", "Re-run all bridge writers (push latest state to all cluster instances)."],
            ["kipi list", "Show all registered projects."],
            ["kipi home", "Print the skeleton path."],
        ],
    )

    add_heading(doc, "13.2 What Propagates and What Does Not", level=2)
    add_para(
        doc,
        "kipi update synchronizes q-system/, .claude/rules/, .claude/agents/, .claude/output-styles/, and the "
        "plugins/ tree. The root CLAUDE.md does NOT propagate. Anything that needs to reach every instance must "
        "live in .claude/rules/ (this is documented as a propagation gotcha in q-system/CLAUDE.md).",
    )

    page_break(doc)

    # ---------- 14. COMPARISON ----------
    add_heading(doc, "14. Comparison: kipi vs memory.md vs Vector Memory", level=1)
    add_para(
        doc,
        "Two common misreadings of kipi: 'this is just memory.md' and 'this is just a memory plugin like MARM.' "
        "Neither is accurate. This section makes the distinctions explicit.",
    )

    add_heading(doc, "14.1 What memory.md Is", level=2)
    add_para(
        doc,
        "memory.md is a single markdown file that Claude reads at session start. It is static text. The user "
        "types it; Claude consumes it. There are no triggers, no schema, no enforcement, and no cross-file "
        "linkage. It does not compound; it does not decay; it does not capture state automatically.",
    )

    add_heading(doc, "14.2 What Vector Memory Adds (and What It Does Not)", level=2)
    add_para(
        doc,
        "A vector memory layer (e.g., a tool that embeds session content into a SQLite + embeddings store) adds "
        "semantic recall: ask 'what did we say about authentication?' and get a relevant chunk back even without "
        "matching keywords. This is genuinely useful for cross-session fuzzy recall.",
    )
    add_para(
        doc,
        "What it does not add: deterministic checks, voice enforcement, scoped multi-instance state, gated "
        "workflows, role-specific rules, follow-up loop tracking, deterministic schedule generation, "
        "council/persona debate, plugin distribution, or any of the behavior layer. Vector memory is one "
        "primitive. Kipi is a system that could use vector memory as one of many stores.",
    )

    add_heading(doc, "14.3 Side-by-Side", level=2)
    add_table_rows(
        doc,
        ["Capability", "memory.md", "Vector memory", "kipi"],
        [
            ["Persistent text loaded at session start", "yes", "yes", "yes (multiple layers)"],
            ["Semantic / fuzzy recall", "no", "yes", "no (deterministic by design)"],
            ["Auto-capture from pasted content", "no", "partial (if explicit)", "yes (auto-detection rules)"],
            ["Conditional load based on context", "no", "no", "yes (rule path globs)"],
            ["Cross-instance state sharing", "no", "no", "yes (bridge protocol)"],
            ["Voice enforcement against output", "no", "no", "yes (skill + lint + rule)"],
            ["Deterministic scorers (real math, no LLM)", "no", "no", "yes (~10 scorers and validators)"],
            ["Append-only event database", "no", "no", "yes (metrics_store SQLite)"],
            ["Schema-validated bus between agents", "no", "no", "yes (35 schemas)"],
            ["Gated multi-step workflows (PRD, issue)", "no", "no", "yes (state machines, scope hooks)"],
            ["Skill / behavior layer", "no", "no", "yes (10+ skills, 19 rules, 5 agents, 6 hooks)"],
            ["Multi-project distribution", "no", "no", "yes (kipi CLI, propagation)"],
        ],
    )

    add_heading(doc, "14.4 The Right Mental Model", level=2)
    add_para(
        doc,
        "Claude Code is the runtime. memory.md is one of many files Claude Code can load. A vector memory layer "
        "is one tool Claude Code can call. Kipi is the operating system that decides which files to load, which "
        "tools to call, which workflows to enforce, how to behave for the specific user, and how to coordinate "
        "across multiple businesses. The runtime executes; the OS decides.",
    )

    page_break(doc)

    # ---------- 15. APPENDIX A ----------
    add_heading(doc, "15. Appendix A - Complete MCP Tool Catalog", level=1)
    add_para(
        doc,
        "The complete inventory of MCP tools registered by the kipi-mcp server, in the order they appear in "
        "server.py:",
    )
    tools_table = [
        ["kipi_suggest_instance_name", "Instance mgmt", "Suggest a unique instance name for a company."],
        ["kipi_set_instance_name", "Instance mgmt", "Set the current instance name."],
        ["kipi_validate", "Validation", "Run the multi-phase validation harness (default phase 5)."],
        ["log_init", "Morning log", "Initialize a daily morning log."],
        ["log_step", "Morning log", "Record a step result (status, result, error)."],
        ["log_add_card", "Morning log", "Add a draft card (target, draft_text, URL)."],
        ["log_deliver_cards", "Morning log", "Mark cards as delivered to founder."],
        ["log_gate_check", "Morning log", "Record a gate check (gate_step, all_prior_done, missing)."],
        ["log_checksum", "Morning log", "Append a checksum (phase, key, value)."],
        ["log_verify", "Morning log", "Record a verification (claim, source, verified, result)."],
        ["loop_open", "Loops", "Open a follow-up loop with target, context, optional Notion id."],
        ["loop_close", "Loops", "Close a loop with reason and closed_by."],
        ["loop_force_close", "Loops", "Force-close a loop with an action."],
        ["loop_escalate", "Loops", "Escalate all loops past their threshold."],
        ["loop_touch", "Loops", "Refresh a loop's last-touched timestamp."],
        ["loop_prune", "Loops", "Prune loops older than N days."],
        ["kipi_create_template", "Content", "Create a content template from a name."],
        ["kipi_build_schedule", "Content", "Build daily schedule HTML from a JSON file."],
        ["kipi_backup", "Portability", "Backup to optional output_path."],
        ["kipi_export", "Portability", "Export to a path."],
        ["kipi_import", "Portability", "Import from an archive (dry-run by default)."],
        ["kipi_verify_schedule", "Bus + verify", "Verify a schedule JSON file for a day."],
        ["kipi_verify_bus", "Bus + verify", "Verify all bus payloads for a date and phase."],
        ["kipi_verify_orchestrator", "Bus + verify", "Verify orchestrator state for date and phase."],
        ["kipi_bus_to_log", "Bus + verify", "Pipe bus payloads to log for a date."],
        ["kipi_scan_draft", "Validation", "Scan a draft JSON file for issues."],
        ["kipi_audit_morning", "Bus + verify", "Audit a morning log file."],
        ["kipi_init_db", "Database", "Initialize the metrics SQLite database."],
        ["kipi_insert_content_metrics", "Database", "Insert content metrics rows."],
        ["kipi_insert_behavioral_signals", "Database", "Insert behavioral signal rows."],
        ["kipi_insert_outreach", "Database", "Insert outreach event rows."],
        ["kipi_insert_copy_edit", "Database", "Insert copy edit rows."],
        ["kipi_query", "Database", "Query a defined query_type with day window."],
        ["kipi_daily_metrics", "Database", "Daily metrics for the founder dashboard."],
        ["kipi_monthly_learnings", "Database", "Monthly aggregated learnings."],
        ["kipi_voice_lint", "Validation", "Lint text against the founder voice rules."],
        ["kipi_validate_schedule", "Validation", "Validate schedule sections for a day."],
        ["kipi_validate_ad_copy", "Validation", "Validate ad headlines and descriptions per platform."],
        ["kipi_seo_check", "Validation", "Run SEO checks against page content."],
        ["kipi_validate_cold_email", "Validation", "Validate a cold email subject and body."],
        ["kipi_copy_edit_lint", "Validation", "Run copy-edit lint on prose."],
        ["kipi_linkedin_gate", "LinkedIn", "Gate a LinkedIn draft against cadence rules."],
        ["kipi_log_linkedin_activity", "LinkedIn", "Log a LinkedIn activity (post, comment, DM)."],
        ["kipi_linkedin_cadence_check", "LinkedIn", "Check today's LinkedIn cadence."],
        ["kipi_score_lead", "Scoring", "Score a lead from attributes and signals."],
        ["kipi_ab_test_calc", "Scoring", "A/B test math: baseline, MDE, traffic, variants."],
        ["kipi_churn_health_score", "Scoring", "Customer churn health score from signals."],
        ["kipi_cancel_flow_offer", "Scoring", "Choose a cancel-flow offer from reason and MRR."],
        ["kipi_crack_detect", "Scoring", "Detect cracks in pipeline from contacts and loops."],
        ["kipi_generate_schema", "Content", "Generate JSON-LD structured data for a page type."],
        ["kipi_harvest", "Harvest", "Run a harvest job."],
        ["kipi_store_harvest", "Harvest", "Store harvested records by source."],
        ["kipi_get_harvest", "Harvest", "Get stored harvest data."],
        ["kipi_harvest_status", "Harvest", "Harvest status by run id."],
        ["kipi_harvest_summary", "Harvest", "Harvest summary by run id."],
        ["kipi_harvest_cleanup", "Harvest", "Clean up harvests older than N days."],
        ["kipi_approve_apify_budget", "Harvest", "Approve an extra Apify budget for a month."],
        ["kipi_harvest_health", "Harvest", "Harvest pipeline health."],
        ["kipi_queue_notion_write", "Notion", "Queue a Notion write with source agent."],
        ["kipi_get_notion_queue", "Notion", "Get the pending Notion write queue."],
        ["kipi_log_agent_metric", "Metrics", "Log an agent metric event."],
        ["kipi_agent_metrics", "Metrics", "Aggregate agent metrics for N days."],
        ["kipi_session_handoff", "Continuity", "Record session handoff (phases completed, notes)."],
        ["kipi_preflight", "Continuity", "Preflight checks for tool availability."],
        ["kipi_session_bootstrap", "Continuity", "Bootstrap a session: read paths, status, instances."],
        ["kipi_canonical_digest", "Continuity", "Build canonical digest from canonical/."],
        ["kipi_morning_init", "Continuity", "Initialize morning routine with an energy level."],
        ["kipi_gate_check", "Bus + verify", "Phase gate check for a date."],
        ["kipi_deliverables_check", "Bus + verify", "Check deliverables for a date."],
    ]
    add_table_rows(doc, ["Tool", "Group", "Purpose"], tools_table)

    page_break(doc)

    # ---------- 16. APPENDIX B ----------
    add_heading(doc, "16. Appendix B - File Inventory", level=1)
    add_para(
        doc,
        "Per-plugin file counts and the location of the major Python modules in kipi-mcp.",
    )
    add_heading(doc, "16.1 Plugin Manifest Summary", level=2)
    add_table_rows(
        doc,
        ["Plugin", "Manifest", "MCP server", "Commands", "Hooks"],
        [
            ["kipi-core", "plugins/kipi-core/.claude-plugin/plugin.json", "plugins/kipi-core/kipi-mcp/src/kipi_mcp/server.py (1694 lines)", "wiring-check (1)", "-"],
            ["kipi-ops", "plugins/kipi-ops/.claude-plugin/plugin.json", "-", "-", "-"],
            ["kipi-design", "plugins/kipi-design/.claude-plugin/plugin.json", "-", "-", "-"],
            ["kipi-dsse", "plugins/kipi-dsse/.claude-plugin/plugin.json", "-", "issue-start, issue-approve, issue-verify, issue-review, issue-amend, issue-closeout (6)", "scope_hook.py, stop_gate.py"],
            ["prd-os", "plugins/prd-os/.claude-plugin/plugin.json", "-", "prd-start, prd-review, prd-approve, prd-triage, prd-split, prd-archive (6)", "scope_hook.py, stop_gate.py"],
            ["kipi-notebooklm", "plugins/kipi-notebooklm/.claude-plugin/plugin.json", "(MCP wrapper via Patchright)", "-", "-"],
        ],
    )

    add_heading(doc, "16.2 kipi-mcp Python Modules", level=2)
    add_table_rows(
        doc,
        ["Module", "Lines", "Responsibility"],
        [
            ["server.py", "1694", "MCP server entrypoint, all tool registrations."],
            ["harvest_store.py", "705", "Harvest CRUD over SQLite."],
            ["linter.py", "600", "Voice lint and copy-edit lint logic."],
            ["morning_init.py", "544", "Morning routine initialization."],
            ["metrics_store.py", "452", "SQLite metrics store schema and queries."],
            ["validator.py", "420", "Validators for ad copy, cold email, schedules."],
            ["harvest_orchestrator.py", "370", "Harvest job orchestration."],
            ["loop_tracker.py", "326", "Follow-up loop CRUD and escalation."],
            ["scorer.py", "313", "Deterministic scoring engines."],
            ["source_registry.py", "263", "External source registry."],
            ["schema_gen.py", "238", "JSON-LD generation."],
            ["paths.py", "230", "Path resolution."],
            ["backup.py", "182", "Backup, export, import."],
            ["schedule_verifier.py", "173", "Schedule integrity checks."],
            ["draft_scanner.py", "157", "Draft scan logic."],
            ["bus_verifier.py", "151", "Bus payload schema verification."],
            ["morning_auditor.py", "145", "Morning log audit."],
            ["orchestrator_verifier.py", "123", "Orchestrator state verification."],
            ["step_logger.py", "113", "Step ledger logic."],
            ["bus_bridge.py", "103", "Bus-to-log piping."],
            ["git_ops.py", "81", "Git status helpers."],
            ["template_manager.py", "73", "Template creation."],
            ["registry.py", "70", "Source registry helpers."],
        ],
    )

    add_heading(doc, "16.3 Hooks in q-system/hooks", level=2)
    add_code(
        doc,
        "auto-commit.py        # Stop hook (async) - commit dirty files\n"
        "auto-update.sh        # Triggered update routine\n"
        "git-health-check.sh   # SessionStart - fleet status banner\n"
        "post-compact.sh       # PostCompact - reinject mode/loops/voice\n"
        "session-context.sh    # Shared session context helpers\n"
        "session-start.py      # SessionStart - memory + paths\n"
        "statusline.sh         # Status line script\n"
        "stop-logger.sh        # Stop hook (async) - session metadata log",
    )

    add_heading(doc, "16.4 Bus Schemas in agent-pipeline/schemas", level=2)
    add_code(
        doc,
        "_bus-envelope.schema.json           bootstrap.schema.json\n"
        "calendar.schema.json                canonical-digest.schema.json\n"
        "client-deliverables.schema.json     collection-gate.schema.json\n"
        "compliance.schema.json              connection-mining.schema.json\n"
        "content-intel.schema.json           copy-diffs.schema.json\n"
        "crm.schema.json                     founder-brand-post.schema.json\n"
        "gmail.schema.json                   hitlist.schema.json\n"
        "leads.schema.json                   linkedin-dms.schema.json\n"
        "linkedin-posts.schema.json          loop-review.schema.json\n"
        "marketing-health.schema.json        meeting-prep.schema.json\n"
        "notion.schema.json                  pipeline-followup.schema.json\n"
        "positioning.schema.json             post-visuals.schema.json\n"
        "preflight.schema.json               prospect-pipeline.schema.json\n"
        "publish-reconciliation.schema.json  schedule-data.schema.json\n"
        "signals.schema.json                 sycophancy-audit.schema.json\n"
        "temperature.schema.json             value-routing.schema.json\n"
        "vc-pipeline.schema.json             warm-intros.schema.json\n"
        "x-activity.schema.json",
    )

    # ---------- SAVE ----------
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size} bytes")


if __name__ == "__main__":
    build()
