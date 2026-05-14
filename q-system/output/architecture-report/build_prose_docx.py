#!/usr/bin/env python3
"""Build the kipi system prose explainer.

Pure narrative. No tables. No diagrams. Just an explanation of how the system
is structured, contrasted against the default Claude Code experience.
"""

from __future__ import annotations

import datetime as dt
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

HERE = Path(__file__).parent
OUTPUT = HERE / "kipi-system-prose-explainer-2026-05-13.docx"


def add_heading(doc, text, level=1, color="1e3a8a"):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)


def para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.35


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


SECTIONS = [
    (
        "1. The Default Claude Code",
        [
            "Claude Code, out of the box, is a runtime. You launch the CLI in a project directory and it gives you a conversation with a Claude model. The model can read and edit files. It can run shell commands. It can call any tool exposed by an MCP server you have configured. It has a context window of around one million tokens, which it fills with whatever it needs to answer the current question.",
            "There are a few seams the runtime exposes for customization. A CLAUDE.md file at the project root gets loaded automatically. Plugins can be installed to add commands, skills, and MCP servers. Hooks can be wired into settings.json to run shell or Python scripts at specific lifecycle events. Skills, when defined inside plugins, can be invoked explicitly by name or auto-loaded by description matching.",
            "This is enough to be useful. Most people who use Claude Code seriously stop here. They write a few hundred lines of CLAUDE.md describing the project, install a handful of MCP servers for the tools they need, and let the model handle the rest of the cognitive load each session. The model is intelligent. The runtime is fast. The context window is large enough to hold most of what matters for most tasks.",
            "The unstated assumption in this setup is that the model is the place where work happens, and the file system is just where artifacts get saved. The conversation is the operating space. Files are output. memory.md, if it exists, is a context primer that the model consumes at the top of each session.",
            "For coding tasks that fit inside a single session, this is fine. For running a business across hundreds of sessions, across multiple lines of work, across years, this collapses in three predictable ways.",
        ],
    ),
    (
        "2. Where the Default Stops Being Enough",
        [
            "The first failure is context loss. A session ends. The next session starts fresh. Anything that was understood, decided, or built during the previous session is gone unless it was explicitly written down. memory.md helps, but it is one file. It cannot hold everything important about every project, every relationship, every decision, every open follow-up, every piece of canonical positioning. It overflows quickly. Once it overflows, you start choosing what to keep and what to drop. The dropped items eventually surface as missed commitments, repeated mistakes, or strategy decisions reversed because the original reasoning was lost.",
            "The second failure is voice drift. Modern language models are trained on enormous amounts of generic web text. Their default writing style reflects that. It is balanced, hedged, transition-heavy, and instantly recognizable as AI to anyone who reads a lot of online content. If you publish under your own name with help from an LLM, you have two choices: rewrite every output by hand, or accept that your audience will eventually clock you. There is no middle ground that the runtime gives you out of the box.",
            "The third failure is role bleed. If you run multiple businesses, the runtime does not help you stay in role. You open Claude Code in your consulting repo and ask for help with a client proposal. You open it in your startup repo and ask about a fundraising message. The model is the same model. Your style preferences are the same. Your sense of who the audience is depends entirely on what context the conversation has so far. If something from the consulting work leaks into the startup output, the model will not notice. You will not notice either, until someone external points it out.",
            "These three failures are not problems with Claude. The runtime is doing what runtimes do. The model is doing what models do. The failures emerge because nobody asked the runtime to be an operating system. It was asked to be a coding assistant. When used as a coding assistant, it succeeds. When used as the operations layer for a real, multi-project, multi-year business, the gaps show.",
            "Kipi exists to close those gaps. It does not try to replace the runtime, fork the model, or build a competing tool. It uses Claude Code as the substrate and adds the operating system on top.",
        ],
    ),
    (
        "3. The Inversion: Files as the System, Not the Output",
        [
            "The defining move in the kipi architecture is inverting what people normally do with Claude Code.",
            "In the default setup, the conversation is the workspace. Files are produced as artifacts. The next session reads whatever notes the previous session wrote down, plus CLAUDE.md, plus whatever the user pastes in.",
            "In kipi, the file system is the workspace. The conversation is a transient surface that reads from and writes to a structured persistent substrate. The model is one participant in the system, not the seat of the system. Files are the seat. The model reads what files instruct it to read, writes to specific places under specific conditions, and exits the session.",
            "This sounds like a small distinction. It is not. It changes who owns context, who enforces behavior, and how state survives time.",
            "Concretely, the kipi-system repository is organized into five durable layers that live on disk: canonical files for long-lived knowledge, rule files for behavioral enforcement, skill bundles for invocable transformations, agent definitions for sub-conversations with their own models, and a bus directory for ephemeral structured messages between pipeline stages. The model reads from these layers conditionally. It writes back to specific places when specific operations succeed. The conversation itself is treated as throwaway. If a session ends mid-thought, the relevant state is on disk. Resuming next session is reading the files back in, not reconstructing from memory.",
            "This inversion is what makes everything else possible. Voice enforcement works because the founder voice profile is a file that the model is required to read before generating external content. Multi-instance isolation works because each instance has its own rule files and each cross-instance share goes through a single-writer bridge file. Deterministic scoring works because the scoring function is Python on disk, not a model call. The model can be replaced, upgraded, or downgraded without changing the system, because the system does not live in the model.",
        ],
    ),
    (
        "4. How File Structure Becomes Behavior",
        [
            "The .claude/rules directory is where the inversion becomes visible. Inside it are nineteen markdown files. Each file declares, in its frontmatter, when it should be loaded into the model's context. Some are always loaded. Some are loaded only when a specific file path is being edited. Some are loaded only when the user's prompt matches a certain pattern.",
            "The rules are not code. They are instructions, written in plain English, that the model reads before it responds. But because Claude Code attaches them automatically based on conditions defined in the rule itself, they behave like enforcement. The model cannot reach a state where it would respond without those rules in context, because the runtime injects them deterministically before each turn.",
            "An example. The social-reaction-gate rule fires when the founder shares someone else's content and asks for a comment, reply, quote, or DM. The rule says: before drafting any reaction, extract the poster's claims as a list, show them to the founder, wait for confirmation, only then draft. This sounds like advice. It functions like a gate, because the rule is in context every time the trigger condition matches, and the model is instructed in clear terms not to skip the extraction step.",
            "Another example. The wiring-check rule is the project's definition of done. It says no implementation task is complete until every change is connected end-to-end across plugins, hooks, MCP tools, agents, bus files, canonical, and rules. When the model edits or writes a file, the post-tool-use hook actually executes the wiring-check Python script as a subprocess. The script walks the changes, looks for orphans, and reports any disconnections. The rule and the script together close the loop: the rule tells the model what done means, and the hook enforces the check at the moment edits land.",
            "Another. The voice-enforcement rule fires when the founder asks the model to draft content for external audiences. The rule does not contain the voice itself. It points to the founder-voice skill, which loads twenty-plus writing samples, five archetypes, banned phrase lists, and the AUDHD-aware writing patterns. The model is required to read these before drafting. After drafting, a separate MCP tool, kipi_voice_lint, runs the draft as a regex against banned words and detects anti-AI tells. The tool returns a pass-or-fail with specific violations. The model can iterate, but it cannot ship until the lint passes.",
            "The pattern in all three examples is the same. A rule lives on disk and is auto-loaded when its triggering condition matches. The rule pins behavior verbally. A separate, deterministic mechanism enforces the rule structurally, either through a hook that runs at a defined lifecycle moment, or through an MCP tool that the model must call before it can claim a result. The rule is the language; the script is the muscle.",
        ],
    ),
    (
        "5. The Deterministic Core",
        [
            "The most important technical claim kipi makes is that numeric and rule-bound work should not run through model inference.",
            "Lead scoring is the cleanest example. If you ask a large language model to score a lead between zero and a hundred based on attributes and signals, you will get a number. If you ask it again with the same input, you might get a different number. The model is using everything in its context, its training distribution, and stochastic sampling. The number drifts. It is also unauditable. You cannot show a teammate or a partner the math, because there is no math.",
            "Kipi's kipi_score_lead tool reads weights from a registry, applies them to the attributes and signals provided, and returns a number. The same input always returns the same output. The weights are reviewable. The math is in scorer.py, around three hundred lines of Python that anyone can read. If the founder disagrees with a score, they can change the weights and rerun.",
            "This pattern repeats. A/B test math runs as arithmetic. Churn health scores run as arithmetic against a defined feature set. Schedule validation walks a JSON schema. Voice linting runs regex against banned words and anti-AI tells. Cold-email validation checks subject length, body length, presence of forbidden phrases, and structural requirements. None of these involve a model.",
            "The reason is conflict of interest, more than performance. A model asked to lint its own output for AI tells has every incentive to say it passed. A model asked to score a lead it just enthusiastically described will rate it higher than the deterministic scorer would. The judgments that matter must live outside the model, in functions whose behavior is fixed at the moment they are written. This is not because the model is dumb. It is because the model is being used elsewhere, and conflating the use and the verification corrupts both.",
            "The model still does enormous amounts of work in kipi. It writes prose. It reasons across canonical files. It decides which MCP tool to call. It synthesizes the daily schedule from twenty-something bus payloads. It does the cognitively expensive work, the kind that requires understanding nuance. The deterministic core handles the work that should be cheap, repeatable, and uncontroversial: math, validation, scoring, schema checks.",
            "About sixty-nine MCP tools live in the kipi-core server. The rough split is that around half are deterministic harness wrappers, around a quarter are bookkeeping (open a loop, close a loop, record a log entry, insert a row), and the remaining quarter expose model-facing capabilities (build a schedule from data, generate JSON-LD, queue a Notion write). The deterministic half is what distinguishes kipi from an opinionated set of prompts. Without it, the system would be advice. With it, the system makes claims that can be checked.",
        ],
    ),
    (
        "6. Hooks as Enforcement",
        [
            "Hooks are the third leg of the deterministic stool, after rules and MCP tools. They are subprocess calls that fire at well-defined moments in the Claude Code session lifecycle. Unlike rules, which are advisory text loaded into context, hooks are code that runs whether or not the model is paying attention.",
            "Kipi wires hooks into six lifecycle events. SessionStart fires when a session begins. It runs three scripts in sequence: a git health check that prints the fleet status banner for all eighteen registered repositories, a session-start Python script that loads the memory index and resolves paths, and a markdown pruning script that archives anything past its freshness threshold. By the time the first user prompt is processed, the model already knows which repos are dirty, which loops are open, and what the canonical state is.",
            "UserPromptSubmit fires before the model sees the user's input. PreToolUse fires before any tool call. Both run the token-guard script, which counts consecutive tool calls without user input and forces a stop if the model is running autonomously past a threshold. This is one of the few hard guardrails in the system: it prevents runaway loops on long-running sessions.",
            "PostToolUse fires after every tool call. The matcher restricts it to Edit and Write, the two operations that mutate the filesystem. When it fires, it runs the wiring-check script, which does the end-to-end connection verification described earlier. This is what makes wiring-check more than advice: every file change triggers an audit.",
            "PostCompact fires when the runtime compacts the conversation history to fit the context window. This is the most subtle hook, because compaction can strip out exactly the kind of session state that the model needs to keep behaving correctly. The post-compact script reinjects the current mode, the list of open loops, and the relevant voice reminders, so that when the conversation resumes after compaction, the model has not silently lost its instructions.",
            "Stop fires at session end. It runs two async scripts: auto-commit, which commits any dirty files to git, and stop-logger, which appends structured session metadata to a daily log. Both run async because the session is ending; the user does not need to wait.",
            "Hooks are not where the intelligence lives. They are where the guarantees live. Every time a session starts, the fleet is checked. Every time a file is edited, the wiring is verified. Every time the conversation is compacted, the critical context is reinjected. These guarantees are not subject to the model deciding they are not important right now. They run as code.",
        ],
    ),
    (
        "7. Skills, Rules, and Agents as Cognitive Layers",
        [
            "Three layers shape how the model thinks, and they are not interchangeable.",
            "Rules are auto-loaded markdown files in .claude/rules. They are short, often under a hundred lines, and they instruct the model about what to do or not do in specific situations. The voice-enforcement rule is twelve lines. The social-reaction-gate rule is twenty. The wiring-check rule is fifty. Their cost is small because they are loaded into context unconditionally when their trigger matches, and the context budget for unconditional rules has to stay tight.",
            "Skills are heavyweight bundles. Each skill has a SKILL.md entrypoint and optional references and templates folders. The founder-voice skill, for example, contains the SKILL.md plus a 400-line voice-DNA reference and a 600-line writing-samples reference. Skills are invoked explicitly, by the model calling the Skill tool with the skill name, or by an auto-detection rule that decides a skill should fire based on the user's input. When a skill is loaded, all of its content goes into context. The cost is significant, so skills are not auto-loaded the way rules are. They are activated only when the work justifies the load.",
            "The split between rules and skills mirrors a familiar engineering pattern: keep the lightweight, always-on policies cheap; reserve the expensive, situational content for explicit invocation. Voice enforcement uses both layers. The rule auto-loads and points to the skill. The model, on seeing the rule, calls the skill if and only if it is about to draft external content. The rule alone is too thin to enforce voice; the skill alone is too expensive to always-load. Together, they cover both the trigger and the substance.",
            "Agents are the third layer, and they work differently from rules and skills. An agent is a sub-conversation. It runs in a fresh context with its own model, its own tool allowlist, its own effort level, and its own system prompt. The orchestrator decides when to spawn an agent based on the work it is doing. Five agents live in .claude/agents: preflight, data-ingest, content-reviewer, engagement-hitlist, and synthesizer. They run on different Claude models depending on the cognitive demand of the work. Data ingestion runs on the cheapest fast model because it is structured extraction with no judgment. The synthesizer runs on the highest-effort Opus configuration because building the daily schedule from twenty bus payloads requires reasoning across all of them.",
            "Together, the three layers form a coarse hierarchy. Rules shape every response that meets their trigger. Skills are loaded when a heavy transformation is needed. Agents are spawned when the work is large enough or specialized enough to deserve its own context window and model selection.",
        ],
    ),
    (
        "8. Memory Without Embeddings",
        [
            "Kipi does not use vector embeddings for memory. This is a deliberate choice, and it cuts against the conventional wisdom that says persistent memory for LLMs requires semantic search.",
            "The reason is determinism. Vector memory retrieves chunks by similarity score. The retrieval is probabilistic in the sense that you cannot guarantee which chunks will surface for a given query. The model gets some context back, but the context is not stable. Run the same query an hour later, after a few new memories have been written, and the top-K may differ. For a system whose value comes from being predictable about which files load when, this is a problem.",
            "Kipi memory works in deterministic layers. The first layer is the user CLAUDE.md, which holds the founder's role, preferences, voice, and AUDHD profile. It always loads. The second layer is the project CLAUDE.md and the q-system CLAUDE.md, which always load when working in this project. The third layer is the .claude/rules directory, where rules load conditionally based on their frontmatter. The fourth layer is the canonical directory, where structured knowledge files live. The fifth layer is the working memory directory, where ephemeral per-day or per-session state lives.",
            "Above these five layers sits the auto-memory subsystem, which writes individual named markdown files into a memory directory and maintains an index in MEMORY.md. Each memory file has a name, a description, and a type tag (user, feedback, project, reference). They link to each other via [[name]] syntax. The index is loaded into every session, which means the model knows what memories exist and can decide whether to read a specific one based on the task at hand. The memories themselves are not pre-loaded; they are pulled on demand by reading the relevant file.",
            "This is closer to how a human filing cabinet works than to how a vector store works. The index is the cabinet drawer labels. The individual memory files are the folders inside. The model looks at the drawer labels, decides which folder is relevant, opens it, reads it, and uses what it found. Nothing is fuzzy. Nothing is approximate. The memory either exists at a specific path, with specific contents, or it does not. If it has been deleted or moved, that is observable. If it has been edited since the last session, the changes are visible.",
            "There is a trade-off. Vector memory gives you fuzzy recall across thousands of memories, including memories you forgot you had. File-based memory requires that you (or the system) maintain enough discipline to give memories names and descriptions that can be searched by description matching. The cost of that discipline buys you predictability. For a system that has to behave reliably across hundreds of sessions and multiple businesses, predictability wins.",
            "Could kipi also use vector memory as a supplemental layer? Yes, easily. Nothing in the architecture forbids it. A vector store could be added as an MCP server alongside the existing ones, and the model could consult it when fuzzy recall would help. But it would be a supplement, not the foundation. The foundation has to be the file-based layers, because those are the layers that allow the rest of the system to make guarantees.",
        ],
    ),
    (
        "9. Multi-Instance and the Bridge Protocol",
        [
            "Kipi is built to run multiple businesses in parallel. Today it runs three: a security product company (KTLYST Labs), a consulting practice (4 Points Consulting), and an advocacy initiative (Pure Spectrum). Each is a distinct repository on disk. Each has its own Claude Code sessions, its own rules, its own canonical knowledge, its own voice variants where appropriate, its own stakeholders.",
            "The naive approach to multi-business operations with an LLM is to dump everything into one workspace and hope the model keeps things straight. This fails immediately. The model has no structural reason to keep contexts separate. Anything in the same conversation is fair game for any output.",
            "Kipi separates instances at the filesystem level. Each business is a directory under ~/projects. Each directory has its own .claude/, its own q-system/, its own canonical files, its own rules. Sessions opened in one directory cannot accidentally read files from another, because Claude Code respects working directory boundaries unless the user explicitly grants cross-directory access with the --add-dir flag.",
            "But the KTLYST cluster has five role-specific instances (strategy, product, website, lawyer, personal-brand) that legitimately need to coordinate. The strategy instance needs to know what the product is actually capable of demoing this week. The product instance needs to know how strategy is positioning what the product is. The website instance needs the latest positioning to keep marketing copy aligned. The lawyer instance needs to flag compliance issues that affect strategy positioning.",
            "Cross-instance coordination is handled through a bridge protocol. A shared directory at ~/.ktlyst/bridge holds a small set of JSON files. Each file has exactly one writer and a defined set of readers. Strategy writes canonical-digest.json and market_signal.json. Product writes product_state.json and threat_status_history.json. Lawyer writes legal-flags.json. Website writes website-state.json. No instance is allowed to write to a bridge file it does not own. No instance reads positioning from anywhere except canonical-digest.json.",
            "This eliminates a class of failures common in poorly-coordinated multi-tenant systems. Two instances cannot disagree about positioning, because positioning has one source of truth. Two instances cannot fight over compliance, because compliance flags come from one place. The price is a small amount of ceremony: when strategy makes a positioning change, the bridge writer script has to run to update canonical-digest.json. The kipi CLI exposes this as kipi cluster sync, which re-runs every bridge writer and pushes the latest state to all instances at once.",
            "The cluster topology also generalizes. The same skeleton runs the non-KTLYST instances (consulting, advocacy, and others) without any cluster wiring. They are independent. When the founder wants to introduce coordination later, the bridge protocol is already designed and ready to extend.",
        ],
    ),
    (
        "10. Gated Workflows: PRD-OS and DSSE",
        [
            "Two of the kipi plugins, prd-os and kipi-dsse, implement gated workflows for product changes. They are the most opinionated, most disciplined part of the system, and they exist because the founder has explicit history with both shipping reckless code and approving sloppy designs.",
            "prd-os is the PRD lifecycle. It models the path of an idea through six states: idea, draft, in-review, approved, split (decomposed into issue specs), and archived. Each state transition is a slash command. Each transition has gates. /prd-start refuses to run if an issue is already in-progress. /prd-approve refuses to run if any finding from the previous review is still pending triage. /prd-archive refuses to run unless every accepted finding has a receipt proving it was either implemented, deferred, or rejected with a written reason.",
            "Reviews run through Codex (a separate AI engineer agent). Codex runs both a native review pass and an adversarial pass that explicitly tries to find issues. Findings are streamed to a JSONL file. The founder triages each finding with one of four dispositions: must-fix, optional, deferred, or rejected-with-reason. No finding is allowed to remain unset. This forces a confrontation with every reviewer concern, even the minor ones.",
            "kipi-dsse picks up where prd-os leaves off. Once a PRD is approved and split into issue specs, the issue specs flow through their own lifecycle: open, in-progress, verified, reviewed, closed. The transitions are commands. The discipline is enforced through two hooks: scope_hook, which runs on every Edit and Write and rejects any edit outside the issue's allowed_files list, and stop_gate, which prevents premature session termination if the issue has not reached its verified and reviewed checkpoints.",
            "The strictest part of this is the empty-allowed_files contract. If an issue declares an empty allowed_files list, the system treats this as deny-all except the spec itself. This is the planning state. The model cannot touch implementation while the issue is still being scoped. Only after /issue-approve transitions the state and arms the stop_gate is implementation work allowed. This structurally prevents the failure mode of starting to code before the issue's scope is locked.",
            "These gated workflows are heavyweight. They are explicitly heavyweight on purpose. The founder does not use them for every small change. They exist for the changes that, if shipped wrong, would be expensive to fix. The lightweight path remains available: just edit the file. The gated path is reserved for product and system work that warrants the ceremony.",
        ],
    ),
    (
        "11. What This Buys",
        [
            "The cumulative effect of these layers is a system that survives time and complexity in a way that a stock Claude Code setup does not.",
            "Voice survives. The founder can paste a draft into any kipi instance and get back a version that passes the lint, sounds like the founder, and avoids the anti-AI tells that detection tools and trained readers pick up on. This is not a one-time fix; it works every time, because the enforcement is structural rather than conversational.",
            "Continuity survives. A meeting today produces a debrief that lands in canonical files, updates relationship state, and surfaces relevant context next week when the same person comes up. The founder does not have to remember to do anything. The auto-detection rule catches the transcript paste and runs the workflow.",
            "Isolation survives. Work in the strategy instance does not contaminate work in the product instance, because they are separate filesystems with separate canonical layers and the only cross-talk goes through bridge files with explicit writer-and-reader contracts. The founder can run a strategy session in the morning and a product session in the afternoon, and the afternoon session has no way of leaking strategy details unless the bridge is explicitly read.",
            "Numeric work survives. Lead scores today match lead scores tomorrow when the inputs are the same, because scoring is deterministic Python. A/B math is auditable. Churn signals are reproducible. The founder can defend a number to a co-founder or investor with the actual computation, not a hand-wave about what the AI said.",
            "Discipline survives. PRD changes go through the gated workflow. Issues respect their scope. Reviews produce findings that must be triaged before close. The founder cannot accidentally ship a half-baked refactor under the cover of an unrelated issue.",
            "Each of these would be possible without kipi, in theory. Each would require constant manual intervention from the founder to maintain. The leverage of kipi is that the maintenance is done by the system, not by the founder, because the maintenance is encoded in files and hooks that run whether or not the founder is paying attention.",
        ],
    ),
    (
        "12. What This Costs",
        [
            "Nothing in this design is free. The costs are real and worth naming.",
            "First, there is a learning cost. A new collaborator or future maintainer cannot walk into a kipi instance and understand what is happening from the conversation alone. They have to read the rules, the skills, the agent definitions, the canonical files, and at least skim the kipi-mcp Python modules. The plurality of layers is the source of the system's power and also the source of its onboarding tax.",
            "Second, there is a maintenance cost. Files drift. Schemas evolve. Rules that made sense six months ago become stale. The system has built-in countermeasures for this (memory-freshness rules, md-prune scripts, the kipi check command that runs validation across the harness), but countermeasures only catch what they are designed to catch. Some drift is invisible until it bites.",
            "Third, there is a propagation cost. Multi-instance distribution means that an improvement to the skeleton has to be rolled out across every registered instance. The kipi update command automates this, but the founder still has to remember to run it, and instances can fall behind if they go untouched for long stretches. Reconciling a stale instance with the latest skeleton is more work than maintaining a fresh one.",
            "Fourth, there is an opinionation cost. Kipi makes strong choices about how things should be done. Voice enforcement is enforced. AUDHD-style output is enforced. Scope hooks reject out-of-bounds edits. These are good choices for the founder. They might be the wrong choices for a different person. The architecture supports being modified, but a future user who disagrees with the core assumptions would have a non-trivial rewrite ahead of them.",
            "Fifth, there is a complexity ceiling. Adding new capabilities means adding new rules, new tools, possibly new agents, and updating the wiring-check coverage so they are not dangling. Each addition is a small amount of work, but the total surface keeps growing. At some point a refactoring round will be necessary. The system has not hit that point yet, but it is in the future.",
            "These costs are accepted because the alternative, doing all the same work manually with each Claude Code session, is worse. The founder has measured this in lost decisions, repeated mistakes, and time spent typing context back in. Kipi pays for itself in continuity. But the costs are real, and a reader evaluating whether to adopt a similar pattern should not skip them.",
        ],
    ),
    (
        "13. The Thesis",
        [
            "The thesis is small and the consequences are large.",
            "Claude Code is a runtime. It executes conversations and tool calls fast and well. It is not, by itself, an operating system. An operating system has policies, services, persistent state, isolation between users, and predictable behavior across reboots. Claude Code provides the machinery; it does not provide the policy.",
            "Kipi is the policy. It sits on the file system, between the runtime and the founder's actual work. It tells the runtime what to load when. It tells the model how to behave for this user, on this project, in this voice. It enforces guarantees through hooks and deterministic Python rather than through prompt instructions alone. It coordinates state across multiple running instances through a single-writer bridge. It treats files as the seat of context and the conversation as a transient surface.",
            "The runtime did not have to be this way. A different team building on top of Claude Code might choose to build their operating system inside the conversation itself, leaning heavily on memory.md and dynamic prompt construction. That choice has different trade-offs, and is reasonable for users whose work fits inside that envelope. For a founder running multiple businesses across multiple years, where context loss, voice drift, role bleed, and numeric drift are all unacceptable, kipi's choice to put the operating system on disk is the one that survives.",
            "The distinction is the difference between an AI assistant that helps you and an operating system you live inside. The former is what most people have. The latter is what kipi is. They are not the same kind of thing, and they should not be evaluated against each other as if they were.",
        ],
    ),
]


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Kipi System")
    tr.font.size = Pt(36)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("How It Is Built and Why It Works")
    sr.font.size = Pt(18)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    doc.add_paragraph()

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr2 = sub2.add_run(
        "A prose explanation of the kipi architecture, contrasted against\n"
        "the default Claude Code experience. No diagrams. No tables.\n"
        "Just how the system uses the file system and a deterministic core\n"
        "to do what it does."
    )
    sr2.font.size = Pt(12)
    sr2.font.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        f"Compiled {dt.date(2026, 5, 13).isoformat()}\n"
        "Author: Assaf Kipnis"
    )
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    page_break(doc)

    # Sections
    for title, paragraphs in SECTIONS:
        add_heading(doc, title, level=1)
        for p_text in paragraphs:
            para(doc, p_text)
        doc.add_paragraph()

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size} bytes")

    total_words = sum(
        len(p.split()) for _, paragraphs in SECTIONS for p in paragraphs
    )
    print(f"Word count (body prose): {total_words}")


if __name__ == "__main__":
    build()
