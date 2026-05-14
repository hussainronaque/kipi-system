#!/usr/bin/env python3
"""Build the kipi novelty research docx.

Honest research finding: is kipi a breakthrough, a novel combination, or nothing new?
Compares against actual projects found on GitHub and discussed on reddit.
"""

from __future__ import annotations

import datetime as dt
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
OUTPUT = HERE / "kipi-novelty-research-2026-05-13.docx"


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color="1e3a8a"):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)


def para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.35


def add_table_rows(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "1e40af")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ----- TITLE -----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("Is Kipi Novel?")
    tr.font.size = Pt(32)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Reddit + GitHub research on what exists, what is comparable, and what is genuinely new")
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = desc.add_run(
        "An honest comparison of kipi-system against the most prominent Claude Code\n"
        "frameworks, memory systems, plugin collections, multi-agent orchestrators,\n"
        "voice/style enforcers, and gated workflow tools available in mid-2026."
    )
    dr.font.size = Pt(11)
    dr.font.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        f"Compiled {dt.date(2026, 5, 13).isoformat()}\n"
        "Author: Assaf Kipnis"
    )
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    page_break(doc)

    # ----- 1. METHOD -----
    add_heading(doc, "1. Method", level=1)
    para(
        doc,
        "This research was conducted on 2026-05-13 using web search across reddit, github, medium, "
        "developer blogs, and Hacker News. Search terms were targeted to surface projects in the same "
        "categories as kipi: Claude Code memory systems, multi-instance frameworks, plugin architectures, "
        "voice and anti-AI enforcement, gated workflows with adversarial review, multi-agent orchestrators, "
        "and neurodivergent-focused founder operating systems."
    )
    para(
        doc,
        "Each project surfaced was evaluated for direct overlap with kipi's capabilities. The findings are "
        "organized by capability dimension (memory, multi-instance, plugins, voice enforcement, gated "
        "workflows, neurodivergent design, founder framing). Each dimension closes with a verdict on whether "
        "kipi is novel, derivative, or one of many comparable implementations."
    )
    para(
        doc,
        "The verdict is delivered honestly. Where competitors are better than kipi, this document says so. "
        "Where kipi is genuinely distinct, this document says that. The point is not to advocate for kipi; "
        "the point is to understand its actual position in the ecosystem as it stands in mid-2026."
    )

    page_break(doc)

    # ----- 2. DIM: MEMORY -----
    add_heading(doc, "2. Dimension: Persistent Memory Across Sessions", level=1)
    para(
        doc,
        "Persistent memory for Claude Code is a mature category in 2026. The dominant player is claude-mem, "
        "a project at roughly 46,100 stars on GitHub, with 1,840 commits, 109 contributors, and 259 releases "
        "in approximately seven months. It uses five hooks (SessionStart, PostToolUse, Stop, UserPromptSubmit, "
        "SessionEnd) to record observations automatically, processes raw observations through Claude's "
        "agent-sdk into semantic summaries stored in SQLite with FTS5 full-text search, and adds a Chroma "
        "vector database for semantic search alongside keyword matching. It installs in one command (npx "
        "claude-mem install) and works across multiple agent systems including Claude Code, OpenClaw, Codex, "
        "Gemini, Hermes, Copilot, and OpenCode."
    )
    para(
        doc,
        "Other entrants in the same space include the native Anthropic memory tool (file-based directory "
        "abstraction in the Claude API), CLAUDE.md plus auto-memory (built into Claude Code), pmem (local "
        "RAG with Ollama + ChromaDB), and various community tools at smaller scales."
    )
    para(
        doc,
        "Kipi's memory architecture is intentionally file-based with no vector embeddings. It uses CLAUDE.md, "
        "auto-memory markdown files with [[link]] cross-references, conditionally loaded rules from "
        ".claude/rules, canonical knowledge files in q-system/canonical, working bus payloads in JSON, "
        "single-writer cross-instance bridge files, and an append-only SQLite metrics store."
    )
    para(
        doc,
        "Verdict: NOT NOVEL. claude-mem is more powerful for pure memory use cases because semantic vector "
        "recall is genuinely useful for finding relevant past observations that don't match keywords. Kipi's "
        "choice of file-based memory is a different trade-off (predictability over fuzzy recall), and that "
        "trade-off makes sense for an operating system, but \"file-based memory for Claude Code\" is not a "
        "unique pattern. Several projects do this. claude-mem itself is the closer match to what most users "
        "mean when they say \"persistent memory for Claude Code,\" and it is materially more developed than "
        "kipi's memory layer."
    )

    page_break(doc)

    # ----- 3. DIM: MULTI-INSTANCE -----
    add_heading(doc, "3. Dimension: Multi-Instance and Multi-Repo Coordination", level=1)
    para(
        doc,
        "Multi-repo and multi-instance Claude Code is a known and active pattern in 2026. Several projects "
        "implement this category at different levels of sophistication."
    )
    para(
        doc,
        "BASE (Builder's Automated State Engine, GitHub: ChristopherKahler/base) is positioned as an \"AI "
        "builder operating system\" that auto-registers projects into a workspace.json, maintains an "
        "operator.json for persistent founder identity, and runs lightweight Python hooks that output "
        "compact XML summaries. It composes with CARL (Context Augmentation & Reinforcement Layer) for "
        "rules and decisions. The BASE+CARL pair maps closely to kipi's separation between project state "
        "and rules."
    )
    para(
        doc,
        "SuperClaude has automated GitHub Actions for distributing a framework into a plugin format using "
        "rsync (very similar to kipi update). Multi-Repo Orchestrator is a skill for tracking repository "
        "dependency graphs and detecting breaking changes across service boundaries. The \"Virtual Monorepo\" "
        "pattern uses workspace files to assemble distributed systems into a single context. Git worktrees "
        "and Dev Containers are well-documented patterns for parallel Claude Code sessions across repos."
    )
    para(
        doc,
        "Kipi's bridge protocol with single-writer JSON files at ~/.ktlyst/bridge/ is a specific design "
        "choice. The single-writer constraint is opinionated; not every multi-instance framework enforces "
        "writer-and-reader contracts. But the broader pattern (shared state files between cooperating "
        "Claude Code instances) is well-precedented."
    )
    para(
        doc,
        "Verdict: NOT NOVEL. Multi-repo Claude Code is now a standard pattern with multiple well-developed "
        "implementations. BASE is the closest direct analog in spirit and structure. Kipi's bridge protocol "
        "is more disciplined than most (single-writer enforcement, defined reader sets, authority rules), "
        "but the core idea of cross-instance state sharing is broadly available."
    )

    page_break(doc)

    # ----- 4. DIM: PLUGINS -----
    add_heading(doc, "4. Dimension: Plugin Architecture with Deterministic Tools", level=1)
    para(
        doc,
        "Plugin architecture for Claude Code is mature in 2026 and there are projects at much larger scale "
        "than kipi."
    )
    para(
        doc,
        "wshobson/agents combines 185 specialized AI agents, 16 multi-agent workflow orchestrators, 153 "
        "agent skills, and 100 commands organized into 80 focused single-purpose plugins. Each installed "
        "plugin loads only its specific agents, commands, and skills into context. The architecture uses "
        "Sonnet + Haiku orchestration with model selection per task complexity. This is materially larger "
        "than kipi's six plugins."
    )
    para(
        doc,
        "claude-flow (ruvnet/ruflo) provides an orchestration layer with MCP server, router, 27 hooks, "
        "swarm coordination, 100+ specialized agents, hybrid memory (AgentDB with vector search), and "
        "support for multiple LLM providers (Claude, GPT, Gemini, Cohere, Ollama). It is positioned as "
        "\"the leading agent orchestration platform for Claude.\""
    )
    para(
        doc,
        "Bernstein is a deterministic orchestrator that spawns parallel AI coding agents and verifies with "
        "tests, with auto-commits and zero LLM tokens on coordination. This is the closest analog to kipi's "
        "deterministic-checks-via-Python philosophy in the orchestration space."
    )
    para(
        doc,
        "Kipi's six plugins (kipi-core, kipi-ops, kipi-design, kipi-dsse, prd-os, kipi-notebooklm) with "
        "approximately 85 deterministic Python tools is modest by 2026 standards. The tools are highly "
        "specific (voice lint, lead scorer, schedule builder, A/B math, churn signals) and oriented toward "
        "founder operations rather than software development."
    )
    para(
        doc,
        "Verdict: NOT NOVEL on architecture. wshobson/agents and claude-flow are bigger. NOVEL in scope: "
        "most large plugin collections target coders. Kipi's tools target founders running businesses, with "
        "specific outputs (lead scores, schedule HTML, voice lints) that do not appear in coder-focused "
        "collections. So the architecture is derivative; the tool selection is unusual."
    )

    page_break(doc)

    # ----- 5. DIM: VOICE -----
    add_heading(doc, "5. Dimension: Voice and Anti-AI Style Enforcement", level=1)
    para(
        doc,
        "Voice and anti-AI enforcement is a small but growing category in 2026."
    )
    para(
        doc,
        "yzhao062/agent-style is a curated set of 21 writing rules formatted for AI coding and writing "
        "agents. It splits rules into deterministic (mechanical checks) and semantic (host-evaluated) "
        "categories. It is integrated with the anywhere-agents ecosystem as a default rule pack."
    )
    para(
        doc,
        "conorbronsdon/avoid-ai-writing is a skill that audits and rewrites content to remove AI writing "
        "patterns, usable with Claude Code, OpenClaw, and Hermes."
    )
    para(
        doc,
        "These are both lightweight projects that focus specifically on voice. Kipi's voice enforcement is "
        "more integrated: a founder-voice skill with 400-line voice DNA references and 600-line writing "
        "samples, a kipi_voice_lint MCP tool that runs Python regex against banned phrases and AI tells, a "
        "voice-enforcement rule that auto-attaches when external content is being drafted, and the "
        "content-reviewer agent that runs a 4-pass review before content ships."
    )
    para(
        doc,
        "Verdict: PARTIALLY NOVEL. The components (anti-AI lists, banned phrases, style rules) exist "
        "elsewhere as separate tools. The integration (rule auto-loads -> skill auto-invokes -> deterministic "
        "lint runs -> agent reviews) is more end-to-end than the standalone tools. The founder-specific "
        "voice profile (with five archetypes, scar pattern, contrast pattern, question-as-dagger) is unique "
        "to the founder, but the pattern of building such a profile is not unique. Net: derivative on "
        "mechanism, distinctive on integration depth."
    )

    page_break(doc)

    # ----- 6. DIM: GATED -----
    add_heading(doc, "6. Dimension: Gated Workflows with Adversarial AI Review", level=1)
    para(
        doc,
        "Gated PRD workflows with multi-model adversarial review are an established pattern in 2026."
    )
    para(
        doc,
        "zscole/adversarial-spec is a Claude Code plugin that iteratively refines product specifications by "
        "debating between multiple LLMs (GPT, Gemini, Grok, etc.) until all models reach consensus. Claude "
        "drafts, opponent models critique, Claude synthesizes. Models that agree within the first two rounds "
        "are pressed for deeper critiques to prevent false convergence. This is a direct analog to kipi's "
        "prd-os adversarial review pattern."
    )
    para(
        doc,
        "shinpr/claude-code-workflows provides production-ready development workflows with specialized "
        "AI agents (requirements, design, implementation, quality). The /recipe-implement command generates "
        "PRD, ADR, Design Doc with acceptance criteria, and a work plan decomposed into commit-ready tasks. "
        "It includes a Codex review loop with the /codex:adversarial-review command and a Stop hook that "
        "runs targeted Codex review based on Claude's response."
    )
    para(
        doc,
        "anombyte93/prd-taskmaster does PRD generation with taskmaster integration. BMAD PRD Workflow is "
        "another Claude Code skill in the same category."
    )
    para(
        doc,
        "Kipi's prd-os adds state machines (idea, draft, in-review, approved, split, archived), scope hooks "
        "that reject edits outside the active PRD or issue file, mandatory dispositions per finding, signed "
        "receipts between issue approve and close, and concurrent PRD/issue blocking. The discipline is "
        "tighter than most analogs."
    )
    para(
        doc,
        "Verdict: NOT NOVEL on concept. Adversarial review and PRD workflows for Claude Code are a saturated "
        "category. Kipi's state-machine plus scope-hook plus receipt enforcement is more disciplined than "
        "most, but the core pattern exists in at least three other public projects with similar or greater "
        "polish."
    )

    page_break(doc)

    # ----- 7. DIM: AUDHD -----
    add_heading(doc, "7. Dimension: Neurodivergent / AUDHD Founder Operating System", level=1)
    para(
        doc,
        "This is the dimension where kipi appears to be genuinely distinct."
    )
    para(
        doc,
        "Searches for \"neurodivergent ADHD AUDHD AI operating system founder workflow github\" surface "
        "kipi-system itself as the primary result. The other projects in the neurodivergent space are "
        "different categories: Leantime is an open-source project management tool designed for ADHD, "
        "Goblin Tools provides ADHD task breakdown utilities, and various smaller projects focus on "
        "specific neurodivergent productivity needs."
    )
    para(
        doc,
        "None of these are operating systems on top of Claude Code with explicit AUDHD layers as a "
        "first-class design principle. The kipi system embeds AUDHD-aware patterns at multiple layers: "
        "the audhd-executive-function skill applied to all actionable output, the audhd-interaction rule "
        "that prevents pressure language and demand framing, an AUDHD output style that enforces no "
        "emdashes and varied sentence rhythm, energy-tagged tasks (Quick Win / Deep Focus / People / "
        "Admin) with time estimates, RSD-safe feedback patterns, and an explicit dismissal rule (when "
        "the user says \"no\" the topic closes)."
    )
    para(
        doc,
        "Articles and reviews of kipi-system explicitly call out the AUDHD-first design: \"The creator has "
        "AUDHD (ADHD + Autism), and every design decision comes from that.\" This is unusual positioning "
        "in the Claude Code framework space, which is otherwise dominated by neurotypical coder-focused "
        "tools."
    )
    para(
        doc,
        "Verdict: GENUINELY NOVEL. No other Claude Code operating system foregrounds neurodivergent "
        "executive function as a first-class design layer. The closest analogs (Leantime, Goblin Tools) "
        "are not Claude Code projects. The closest Claude Code analogs (BASE, claude-mem, wshobson/agents) "
        "do not have neurodivergent-specific design. This is the dimension where kipi has the strongest "
        "claim to originality."
    )

    page_break(doc)

    # ----- 8. DIM: FOUNDER FRAMING -----
    add_heading(doc, "8. Dimension: Founder-First vs Coder-First Operating System", level=1)
    para(
        doc,
        "Most Claude Code frameworks in 2026 are coder-focused. wshobson/agents covers python-development, "
        "backend-development, security-scanning. claude-flow does swarm coordination for software tasks. "
        "BASE manages PAUL projects (a code-builder framework). claude-mem is targeted at long coding "
        "sessions. The plurality of the ecosystem is shaped around software development workflows."
    )
    para(
        doc,
        "An emerging but small category is the founder-first operating system. The aiadopters.club article "
        "titled \"A Personal Operating System for Founders, Built in 10 Minutes with Claude Code\" "
        "represents this category. The kipi-system itself markets itself as \"A portable founder operating "
        "system for Claude Code\" with morning briefings, conversation debriefs, social engagement, "
        "relationship tracking, content pipeline, and lead sourcing as primary use cases."
    )
    para(
        doc,
        "Onur Polat's Medium piece (\"Claude Code: More Than a Coder, It's Your Personal OS\") and Rick "
        "Hightower's piece (\"Claude Code 2026: The Daily Operating System Top Developers Actually Use\") "
        "argue that Claude Code can be used as a daily operating system but are still developer-framed. The "
        "explicit founder-not-coder framing is less common."
    )
    para(
        doc,
        "Verdict: SEMI-NOVEL. Founder-OS framing for Claude Code exists but is small and emerging. Kipi is "
        "one of a handful of public projects positioned explicitly as a founder operating system rather "
        "than a coder operating system. The framing alone is not unique, but the depth of execution against "
        "founder-specific use cases (relationship tracking, voice for outreach, lead scoring, daily schedule "
        "synthesis, follow-up loop tracking) is unusual in the ecosystem."
    )

    page_break(doc)

    # ----- 9. THE COMBINATION -----
    add_heading(doc, "9. The Combination: Does the Integration Matter?", level=1)
    para(
        doc,
        "Most analyses of \"is X novel\" stop at the individual dimensions. That gives a misleading answer "
        "when the system's value is in the integration. Here is the side-by-side coverage across the major "
        "competitors and kipi:"
    )
    add_table_rows(
        doc,
        ["Capability", "claude-mem", "wshobson/agents", "claude-flow", "BASE+CARL", "adversarial-spec", "kipi"],
        [
            ["Persistent memory", "yes (vector)", "no", "yes (vector)", "yes (file)", "no", "yes (file)"],
            ["Multi-instance coord", "no", "no", "yes", "yes", "no", "yes"],
            ["Plugin architecture", "no", "yes (80 plugins)", "yes", "yes", "no", "yes (6 plugins)"],
            ["Voice/anti-AI lint", "no", "no", "no", "no", "no", "yes"],
            ["Gated PRD workflow", "no", "partial", "partial", "no", "yes", "yes"],
            ["AUDHD design layer", "no", "no", "no", "no", "no", "yes"],
            ["Founder-first framing", "no", "no", "no", "yes", "no", "yes"],
            ["Skeleton propagation", "no", "no", "no", "yes", "no", "yes"],
            ["Bridge protocol (single-writer)", "no", "no", "no", "no", "no", "yes"],
            ["Schema-validated bus", "no", "no", "yes", "no", "no", "yes"],
        ],
    )
    para(
        doc,
        "Reading the table: claude-mem wins on memory, wshobson/agents wins on plugin count, claude-flow "
        "wins on multi-agent sophistication, BASE+CARL wins on operator identity and workspace state, "
        "adversarial-spec wins on PRD review depth. No single project covers all the dimensions kipi "
        "covers."
    )
    para(
        doc,
        "This does not mean kipi is the best at any single dimension. It is materially smaller than "
        "wshobson/agents and claude-flow, materially less developed than claude-mem on memory, and roughly "
        "on par with BASE+CARL on the state-engine pieces. What kipi does that no single competitor does "
        "is integrate all of these dimensions plus the AUDHD layer plus the founder-first framing into "
        "one coherent system."
    )
    para(
        doc,
        "Whether the combination matters depends on the user. For somebody who just needs better memory "
        "in Claude Code, claude-mem is a better answer. For somebody who needs sophisticated multi-agent "
        "orchestration for software work, claude-flow is a better answer. For somebody running multiple "
        "businesses with neurodivergent needs and wants voice enforcement and gated workflows in one "
        "system, there is no off-the-shelf competitor that does what kipi does."
    )

    page_break(doc)

    # ----- 10. VERDICT -----
    add_heading(doc, "10. The Verdict", level=1)
    para(
        doc,
        "Honest answer: kipi is not a breakthrough on any single technical dimension. Every piece of the "
        "architecture exists somewhere else, often in projects that are larger, more polished, or more "
        "established. The dominant memory framework (claude-mem) is materially more powerful for memory. "
        "The dominant plugin collection (wshobson/agents) is materially bigger. The dominant orchestrator "
        "(claude-flow) is more sophisticated. The dominant gated-workflow tools (adversarial-spec, "
        "claude-code-workflows) are at parity or ahead on PRD discipline. The closest spiritual analog "
        "(BASE+CARL) covers the same operating-system framing with comparable polish."
    )
    para(
        doc,
        "What kipi does that nothing else does is integrate at this depth for this specific user. The "
        "AUDHD-first design is unique. The founder-not-coder framing is among a small handful. The "
        "combination of memory, multi-instance, plugins, voice, gated workflows, schema bus, skeleton "
        "propagation, and bridge protocol in one coherent system is not available as a single project "
        "elsewhere."
    )
    para(
        doc,
        "Two ways to read this honestly:"
    )
    para(
        doc,
        "If \"novel\" means \"introduces a new technique nobody has used,\" kipi is not novel. The "
        "techniques are widely available."
    )
    para(
        doc,
        "If \"novel\" means \"integrates known techniques into a coherent operating system for a specific "
        "user that no off-the-shelf tool serves,\" kipi is novel. The integration target (a neurodivergent "
        "founder running multiple businesses with enforced voice, gated workflows, and isolated multi-"
        "instance state) is not addressed by any single competitor, and the integration depth (rules + "
        "skills + agents + hooks + MCP + bus + canonical + bridges + cluster) is not seen as a stack."
    )
    para(
        doc,
        "The right framing is probably: kipi is a thoughtful integration of mature patterns with one "
        "genuinely original layer (AUDHD-first design) and a specific positioning (founder OS, not coder "
        "OS) that is still emerging in the ecosystem. It is not a breakthrough that the field hadn't seen. "
        "It is not nothing new either. It is a competent personal operating system that uses the available "
        "primitives intelligently, with one part that no other public project does."
    )
    para(
        doc,
        "For somebody whose situation matches the founder's situation, kipi is the closest thing on the "
        "market. For somebody whose situation doesn't match, a more specialized tool is probably a "
        "better answer for that specific need."
    )

    page_break(doc)

    # ----- 11. WHAT TO SAY -----
    add_heading(doc, "11. What This Means for Conversations on Reddit and Beyond", level=1)
    para(
        doc,
        "The right framing for skeptical builders asking \"isn't this just X?\" is to acknowledge what's "
        "true and name what's distinct."
    )
    para(
        doc,
        "When somebody says \"this is just memory.md / native Claude / claude-mem\": agree on the substrate "
        "and name the dimensions they're missing (hooks, rules, multi-instance, voice lint, gated workflows). "
        "Specifically name claude-mem if relevant; it is the legitimate memory winner."
    )
    para(
        doc,
        "When somebody says \"this is just claude-flow / wshobson/agents\": agree those are bigger on "
        "their core dimension (orchestration, plugin count) and name what kipi covers that they do not "
        "(AUDHD layer, founder framing, voice enforcement, single-writer bridge protocol)."
    )
    para(
        doc,
        "When somebody says \"why not BASE+CARL\": acknowledge BASE is the closest analog in spirit, and "
        "name the specific dimensions kipi adds (AUDHD design, founder-specific scorers like lead scoring "
        "and churn signals, multi-business bridge protocol)."
    )
    para(
        doc,
        "Resist the temptation to claim kipi is breakthrough technology. It is not. Claim it as a thoughtful "
        "integration with one original layer (AUDHD), which is the honest position. Builders respect honesty; "
        "they will challenge overclaiming and reward calibration."
    )

    page_break(doc)

    # ----- 12. SOURCES -----
    add_heading(doc, "12. Sources", level=1)
    para(doc, "Primary projects evaluated:")
    para(doc, "claude-mem (thedotmack/claude-mem): https://github.com/thedotmack/claude-mem")
    para(doc, "wshobson/agents: https://github.com/wshobson/agents")
    para(doc, "claude-flow / ruflo (ruvnet): https://github.com/ruvnet/claude-flow")
    para(doc, "BASE (ChristopherKahler/base): https://github.com/ChristopherKahler/base")
    para(doc, "seed (ChristopherKahler/seed): https://github.com/ChristopherKahler/seed")
    para(doc, "adversarial-spec (zscole): https://github.com/zscole/adversarial-spec")
    para(doc, "claude-code-workflows (shinpr): https://github.com/shinpr/claude-code-workflows")
    para(doc, "agent-style (yzhao062): https://github.com/yzhao062/agent-style")
    para(doc, "avoid-ai-writing (conorbronsdon): https://github.com/conorbronsdon/avoid-ai-writing")
    para(doc, "open-gitagent / gitagent: https://github.com/open-gitagent/gitagent")
    para(doc, "everything-claude-code (affaan-m): https://github.com/affaan-m/everything-claude-code")
    para(doc, "agent-governance-toolkit (microsoft): https://github.com/microsoft/agent-governance-toolkit")
    para(doc, "Anthropic Claude Code memory docs: https://code.claude.com/docs/en/memory")
    para(doc, "Anthropic memory tool API docs: https://platform.claude.com/docs/en/agents-and-tools/tool-use/memory-tool")
    para(doc, "Latent Space podcast (Anthropic's RAG decision): referenced via SmartScope and Medium articles, May 2025")
    para(doc, "")
    para(doc, "Secondary sources:")
    para(doc, "MindStudio blog comparing Claude Code memory systems")
    para(doc, "SmartScope blog on RAG vs agentic search decision in Claude Code")
    para(doc, "aiadopters.club article on founder personal operating systems")
    para(doc, "Rick Hightower, Onur Polat Medium pieces on Claude Code as personal OS")
    para(doc, "DEV Community: Multi-repo Claude Code patterns, virtual monorepo")
    para(doc, "Bridgers Agency: claude-mem review and alternatives")

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size} bytes")


if __name__ == "__main__":
    build()
