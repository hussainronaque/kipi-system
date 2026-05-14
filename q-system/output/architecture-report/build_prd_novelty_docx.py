#!/usr/bin/env python3
"""Build the kipi PRD-OS novelty research docx."""

from __future__ import annotations

import datetime as dt
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
OUTPUT = HERE / "kipi-prd-os-novelty-2026-05-13.docx"


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color="1e3a8a"):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)


def para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.35


def add_table_rows(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "1e40af")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- TITLE ----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("Is kipi's PRD-OS Novel?")
    tr.font.size = Pt(30)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("A focused review of the gated PRD + Issue workflow space")
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = desc.add_run(
        "An honest evaluation of kipi's PRD-OS plugin (prd-os + kipi-dsse) against\n"
        "the eight or more public Claude Code projects in the same category.\n"
        "Where it lands. What it does that nothing else does. What others do better."
    )
    dr.font.size = Pt(11)
    dr.font.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        f"Compiled {dt.date(2026, 5, 13).isoformat()}\n"
        "Author: Assaf Kipnis"
    )
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    page_break(doc)

    # ---- 1. THE CATEGORY ----
    add_heading(doc, "1. The PRD-Workflow Category in 2026", level=1)
    para(
        doc,
        "Gated PRD workflows with adversarial AI review have become a saturated category in the Claude Code "
        "ecosystem by mid-2026. The pattern is well-defined: Claude (or another primary model) drafts a PRD or "
        "spec; one or more reviewer models critique it; the human triages findings; only after all findings "
        "are addressed does the work advance to implementation. Some systems extend the pattern further into "
        "issue execution with scope enforcement, TDD, or verification protocols."
    )
    para(
        doc,
        "What unites this category is the recognition that letting a single model draft and implement without "
        "review is dangerous. What separates the implementations is how strictly the gates are enforced, how "
        "the adversarial review is structured, and how the workflow integrates with downstream implementation."
    )
    para(
        doc,
        "Public projects identified in this category, in rough order of polish and adoption:"
    )

    page_break(doc)

    # ---- 2. THE COMPETITORS ----
    add_heading(doc, "2. The Public Projects in This Space", level=1)

    add_heading(doc, "2.1 BMAD-METHOD", level=2)
    para(
        doc,
        "BMAD (Build More Architect Dreams) is the most established project in the space, with the longest "
        "history and most polish. It is positioned as a \"Breakthrough Method for Agile AI Driven Development.\" "
        "BMAD's distinguishing feature is dedicated agentic planning roles: an Analyst agent, a PM agent, an "
        "Architect agent, and a Scrum Master agent collaborate with the founder to produce detailed PRDs and "
        "architecture documents. The Scrum Master then transforms these plans into hyper-detailed development "
        "stories with embedded context, which a Dev agent implements. The two-phase split (plan, then "
        "implement) is enforced by treating planning artifacts as a contract."
    )
    para(
        doc,
        "BMAD has both a generic version (bmad-code-org/BMAD-METHOD) and a Claude Code-specific port "
        "(24601/BMAD-AT-CLAUDE). It includes a skills package for Claude Code (aj-geddes/claude-code-bmad-skills) "
        "with auto-detection, memory integration, and slash commands. BMAD is the system most builders in this "
        "category will mention first."
    )

    add_heading(doc, "2.2 pilot-shell (maxritter)", level=2)
    para(
        doc,
        "Pilot Shell positions itself as \"How real engineers run Claude Code: spec-driven planning, enforced "
        "TDD, persistent memory, and quality enforcement on all levels.\" The /spec command replaces Claude "
        "Code's built-in plan mode. TDD enforcement is universal: bugfixes without a failing test cannot ship. "
        "Quality hooks fire automatically across the lifecycle for formatting, linting, type checking, TDD "
        "enforcement, context preservation, and memory capture. Supports Python, TypeScript/JavaScript, and "
        "Go out of the box."
    )
    para(
        doc,
        "Pilot Shell is closer to a discipline framework for working coders than a multi-phase PRD system. "
        "It is the strongest analog for kipi's enforcement-via-hooks philosophy in the pure software "
        "development context."
    )

    add_heading(doc, "2.3 adversarial-spec (zscole)", level=2)
    para(
        doc,
        "Adversarial-spec is a Claude Code plugin that iteratively refines product specifications by debating "
        "between multiple LLMs (GPT, Gemini, Grok, and others) until all models reach consensus. Claude drafts "
        "an initial document. Opponent models critique in parallel. Claude synthesizes the critiques into a "
        "revised spec. Models that agree within the first two rounds are pressed for deeper critiques to "
        "prevent false convergence from rubber-stamping reviewers. When the PRD reaches consensus, the user "
        "can continue directly into a Technical Specification based on the PRD."
    )
    para(
        doc,
        "This is the most sophisticated adversarial review in the category. Where kipi uses Codex as a "
        "single adversary, adversarial-spec uses the entire frontier LLM lineup."
    )

    add_heading(doc, "2.4 claude-code-workflows (shinpr)", level=2)
    para(
        doc,
        "Claude-code-workflows provides production-ready development workflows powered by specialized AI "
        "agents. The /recipe-implement command generates a PRD, an ADR (when applicable), and a Design Doc "
        "with acceptance criteria, along with a work plan decomposed into commit-ready tasks. It includes "
        "a Codex review loop with /codex:review and /codex:adversarial-review commands, plus a Stop hook that "
        "runs targeted Codex review based on Claude's response and blocks the stop if issues are found."
    )
    para(
        doc,
        "This is the closest direct analog to kipi's prd-os in terms of overall shape. It is more mature on "
        "the artifact generation side (PRD plus ADR plus Design Doc) and includes a documented Stop-hook "
        "review loop."
    )

    add_heading(doc, "2.5 prd-taskmaster (anombyte93)", level=2)
    para(
        doc,
        "AI-powered PRD generation for Claude Code with taskmaster integration. Focuses on the PRD-to-tasks "
        "decomposition step. Less elaborate than the other entrants in this list, but represents the broader "
        "trend of plugging Claude Code into existing project-management tools."
    )

    add_heading(doc, "2.6 social-science-claude-scholar (HaipingXu)", level=2)
    para(
        doc,
        "A Claude Code configuration for social science research (Economics and Political Science). 42 skills, "
        "quality gates, adversarial QA, causal inference workflows, replication protocols. Adapts the gated "
        "review pattern to academic research output. Demonstrates that the pattern generalizes beyond software."
    )

    add_heading(doc, "2.7 claude-octopus (nyldn)", level=2)
    para(
        doc,
        "Puts up to 8 AI models on every research, design, or coding task. Positioned as surfacing AI blindspots "
        "before shipping. More about parallel multi-model evaluation than gated workflows specifically, but "
        "directly relevant to the adversarial review part of the category."
    )

    add_heading(doc, "2.8 vinicius91carvalho/.claude", level=2)
    para(
        doc,
        "A portable Claude Workflow System with hooks, agents, skills, and enforcement. Smaller scale than "
        "the others but exemplifies the pattern of packaging discipline as a portable .claude directory."
    )

    page_break(doc)

    # ---- 3. WHAT KIPI DOES DIFFERENTLY ----
    add_heading(doc, "3. What Kipi's prd-os + kipi-dsse Does That I Did Not Find Elsewhere", level=1)
    para(
        doc,
        "After reviewing the eight projects above, here is what kipi's prd-os + kipi-dsse combo does that "
        "did not appear in any single competitor:"
    )

    add_heading(doc, "3.1 Explicit State Machines With Signed Receipts", level=2)
    para(
        doc,
        "Most competitors describe \"phases\" or \"gates\" informally. Kipi has formal state machines. "
        "The PRD lifecycle: idea, draft, in_review, approved, split, archived. The Issue lifecycle: open, "
        "in_progress, verified, reviewed, closed. Each state transition is a slash command. Some transitions "
        "require receipts."
    )
    para(
        doc,
        "Receipts are concrete: \"verified\" is written by /issue-verify when all required_checks pass. "
        "\"reviewed\" is written by /issue-review after Codex native and adversarial review run. "
        "\"findings_triaged\" is written by /issue-closeout after every finding has a disposition. The "
        "receipts are deterministic artifacts on disk, not informal markers. The next state cannot be "
        "entered without the receipt for the previous gate."
    )
    para(
        doc,
        "Other systems have gates. Kipi has gates with proofs."
    )

    add_heading(doc, "3.2 Empty allowed_files = Deny-All During Planning", level=2)
    para(
        doc,
        "While a PRD is in the draft state or an issue is in the open (planning) state, scope_hook.py runs at "
        "PostToolUse and rejects any Edit or Write that touches a file outside the active spec's allowed_files "
        "list. If allowed_files is empty (the default during planning), the model can only edit the spec "
        "itself. Implementation cannot begin while planning is underway."
    )
    para(
        doc,
        "Most competitors rely on agent discipline: \"plan first, then implement\" is an instruction, not an "
        "enforcement. BMAD relies on the Scrum Master agent's role to keep planning and implementation "
        "separate. Pilot Shell uses TDD to keep implementation honest, but does not block implementation "
        "during planning. Kipi structurally prevents the model from coding while the scope is still being "
        "negotiated."
    )

    add_heading(doc, "3.3 Mandatory Dispositions Per Finding", level=2)
    para(
        doc,
        "Codex review writes findings to a JSONL file. Each finding must receive exactly one of four "
        "dispositions before the PRD or issue can advance: must-fix, optional, deferred, or "
        "rejected-with-reason. No finding may stay unset. The triage step is itself a state transition; "
        "/prd-approve and /issue-closeout both refuse to advance if any finding lacks a disposition."
    )
    para(
        doc,
        "Adversarial-spec gets all models to consensus, but does not force the human to confront every "
        "individual finding. Claude-code-workflows runs review, but the triage step is informal. BMAD "
        "produces detailed plans but does not have a finding-disposition contract. Kipi forces a "
        "confrontation with every reviewer concern, even minor ones, before advancing the workflow."
    )

    add_heading(doc, "3.4 Concurrent PRD and Issue Blocking", level=2)
    para(
        doc,
        "/prd-start refuses if an issue is currently in-progress. /issue-start refuses if a PRD is currently "
        "in-review. The system enforces single-channel discipline: the founder cannot have one foot in "
        "planning and one foot in implementation. This is a structural constraint, not a guideline."
    )
    para(
        doc,
        "No competitor I evaluated has this specific blocking behavior. Most assume the user knows what they "
        "are doing. Kipi assumes the user might forget and removes the failure mode."
    )

    add_heading(doc, "3.5 Codex as Pure Reviewer, Claude as Sole Editor", level=2)
    para(
        doc,
        "Codex never edits in kipi. Claude is the sole editor. Codex runs through /codex:review and "
        "/codex:adversarial-review and returns findings for Claude to triage. This is a clean separation of "
        "concerns. The reviewer cannot also be the author. False positives or contested findings get "
        "negotiated through the disposition step, not by allowing the reviewer to make unilateral changes."
    )
    para(
        doc,
        "Most other systems either let the reviewer model propose edits inline or do not specify the editor-"
        "reviewer separation at all. Adversarial-spec is multi-LLM debate; the final synthesis is by Claude, "
        "but the boundary is fuzzier. BMAD has multiple agents that can edit. Kipi formalizes the constraint: "
        "review without write access."
    )

    page_break(doc)

    # ---- 4. SIDE BY SIDE ----
    add_heading(doc, "4. Side-by-Side Comparison", level=1)
    para(
        doc,
        "Mapping kipi against the five most directly comparable projects:"
    )
    add_table_rows(
        doc,
        ["Feature", "BMAD", "pilot-shell", "adversarial-spec", "claude-code-workflows", "kipi prd-os+dsse"],
        [
            ["PRD generation", "yes (multi-agent)", "yes (spec)", "yes", "yes (PRD + ADR + Design)", "yes"],
            ["Adversarial AI review", "informal", "no", "yes (multi-LLM)", "yes (Codex)", "yes (Codex)"],
            ["Multi-LLM consensus", "no", "no", "yes", "no", "no (Codex-only)"],
            ["Formal state machine", "phases", "no", "no", "no", "yes (PRD + Issue)"],
            ["Signed receipts per gate", "no", "no", "no", "no", "yes (verified, reviewed, findings_triaged)"],
            ["Scope hook on edits", "no", "no", "no", "no", "yes (allowed_files)"],
            ["Deny-all during planning", "no", "no", "no", "no", "yes (empty allowed_files)"],
            ["Mandatory finding dispositions", "no", "no", "consensus", "no", "yes (4 dispositions)"],
            ["Concurrent PRD/issue blocking", "no", "no", "no", "no", "yes"],
            ["Editor / reviewer separation", "partial", "TDD enforces", "fuzzy", "Codex-as-reviewer", "yes (Codex never edits)"],
            ["TDD enforcement", "no", "yes (strict)", "no", "partial", "no"],
            ["Story-level decomposition", "yes (Scrum Master agent)", "tasks", "tech-spec", "commit-ready tasks", "/prd-split into issues"],
            ["Project polish + adoption", "highest", "high", "medium", "medium", "personal (single-author)"],
        ],
    )
    para(
        doc,
        "Reading the table: BMAD wins on polish and ecosystem maturity. Pilot Shell wins on TDD enforcement. "
        "Adversarial-spec wins on multi-LLM review sophistication. Claude-code-workflows wins on artifact "
        "coverage (PRD + ADR + Design Doc). Kipi wins on enforcement rigor: state machines, receipts, scope "
        "hooks, mandatory dispositions, concurrent blocking, and editor-reviewer separation are tighter than "
        "anything else in the table."
    )

    page_break(doc)

    # ---- 5. VERDICT ----
    add_heading(doc, "5. The Verdict", level=1)
    para(
        doc,
        "Concept-level: kipi's prd-os and kipi-dsse are not novel. The pattern (PRD with adversarial AI "
        "review, then gated implementation) is mature in 2026 and has multiple established public projects. "
        "BMAD predates and outpolishes anything kipi could claim on the planning side. Pilot Shell is more "
        "developed on the discipline-via-hooks side. Adversarial-spec is more sophisticated on the review "
        "side."
    )
    para(
        doc,
        "Enforcement-level: kipi appears to be the most strict in the category. The combination of state "
        "machines, signed receipts, scope hooks that deny-all during planning, mandatory four-disposition "
        "triage per finding, concurrent PRD/issue blocking, and Codex-as-pure-reviewer is tighter than the "
        "competitors I evaluated. None of the five direct analogs has all five of these constraints."
    )
    para(
        doc,
        "Whether this matters depends on what you are protecting against. If your failure mode is \"AI ships "
        "code without enough review,\" most of these systems will help. If your failure mode is \"I have ADHD "
        "and will silently skip steps if the workflow lets me,\" kipi's enforcement model is more aligned. "
        "The state machines and signed receipts exist specifically to make skipping impossible. The empty-"
        "allowed_files contract exists specifically to prevent the founder from getting lured into coding "
        "while planning is open. The mandatory dispositions exist specifically to prevent rubber-stamping "
        "during triage."
    )
    para(
        doc,
        "Net: kipi's PRD-OS is a strict implementation of an established pattern. Not breakthrough. More "
        "rigorous than most. A reasonable choice for users who want the strictest available enforcement; "
        "an unreasonable choice for users who want the most polished agentic planning experience (BMAD), "
        "the strictest TDD discipline (pilot-shell), or the most sophisticated multi-LLM review "
        "(adversarial-spec)."
    )

    page_break(doc)

    # ---- 6. WHAT TO SAY ----
    add_heading(doc, "6. What This Means for Conversations", level=1)
    para(
        doc,
        "When a builder asks \"isn't this just BMAD?\": acknowledge BMAD has more polish on the planning "
        "agentics, and name what kipi adds: state machines with receipts, scope hook deny-all, mandatory "
        "finding dispositions, concurrent blocking. The shape is shared. The enforcement is stricter."
    )
    para(
        doc,
        "When a builder asks \"isn't this just pilot-shell\": acknowledge pilot-shell's TDD discipline is "
        "stronger and more developer-focused, and name that kipi's PRD-OS is upstream of implementation, "
        "covering the planning-gate space that pilot-shell does not focus on."
    )
    para(
        doc,
        "When a builder asks \"isn't this just adversarial-spec\": acknowledge adversarial-spec's multi-LLM "
        "consensus is more sophisticated than kipi's Codex-only review, and name what kipi adds downstream: "
        "the scope hook, the issue execution receipts, the disposition contract. Adversarial-spec stops at "
        "PRD-to-tech-spec; kipi continues into gated implementation."
    )
    para(
        doc,
        "Resist the temptation to claim originality on the PRD-with-adversarial-review pattern itself. The "
        "pattern is widely adopted. The claim worth defending is enforcement rigor, not invention."
    )

    page_break(doc)

    # ---- 7. SOURCES ----
    add_heading(doc, "7. Sources", level=1)
    para(doc, "BMAD-METHOD (bmad-code-org): https://github.com/bmad-code-org/BMAD-METHOD")
    para(doc, "BMAD-AT-CLAUDE (24601): https://github.com/24601/BMAD-AT-CLAUDE")
    para(doc, "claude-code-bmad-skills (aj-geddes): https://github.com/aj-geddes/claude-code-bmad-skills")
    para(doc, "pilot-shell (maxritter): https://github.com/maxritter/pilot-shell")
    para(doc, "pilot-shell docs: https://pilot-shell.com/")
    para(doc, "adversarial-spec (zscole): https://github.com/zscole/adversarial-spec")
    para(doc, "claude-code-workflows (shinpr): https://github.com/shinpr/claude-code-workflows")
    para(doc, "prd-taskmaster (anombyte93): https://github.com/anombyte93/prd-taskmaster")
    para(doc, "social-science-claude-scholar (HaipingXu): https://github.com/HaipingXu/social-science-claude-scholar")
    para(doc, "claude-octopus (nyldn): https://github.com/nyldn/claude-octopus")
    para(doc, ".claude (vinicius91carvalho): https://github.com/vinicius91carvalho/.claude")
    para(doc, "everything-claude-code (affaan-m): https://github.com/affaan-m/everything-claude-code")
    para(doc, "codex-plugin-cc (openai): https://github.com/openai/codex-plugin-cc")
    para(doc, "Anthropic Claude Code hooks docs: https://code.claude.com/docs/en/hooks-guide")
    para(doc, "")
    para(doc, "Secondary sources:")
    para(doc, "BMAD method documentation: https://docs.bmad-method.org/")
    para(doc, "SmartScope blog on automating Claude × Codex review loop")
    para(doc, "BMAD framework analysis: pasqualepillitteri.it")
    para(doc, "Medium articles on BMAD method with Cursor, Codex, Claude")

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size} bytes")


if __name__ == "__main__":
    build()
