#!/usr/bin/env python3
"""Build the Claude Code + kipi vs Claude.ai comparison docx.

Side-by-side analysis. Few tables for comparison sections, prose elsewhere.
"""

from __future__ import annotations

import datetime as dt
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
OUTPUT = HERE / "kipi-vs-claude-ai-comparison-2026-05-13.docx"


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color="1e3a8a"):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)


def para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.35


def add_table_rows(doc, headers, rows, header_color="1e40af"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, header_color)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ----- TITLE PAGE -----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("Claude Code + kipi vs Claude.ai")
    tr.font.size = Pt(32)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Two different products. Two different jobs.")
    sr.font.size = Pt(16)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = desc.add_run(
        "An honest side-by-side comparison of how the web product (Claude.ai)\n"
        "works versus how Claude Code with the kipi operating system works,\n"
        "based on the actual Anthropic documentation and the kipi codebase.\n\n"
        "Both products are useful. They are not the same kind of tool."
    )
    dr.font.size = Pt(11)
    dr.font.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        f"Compiled {dt.date(2026, 5, 13).isoformat()}\n"
        "Author: Assaf Kipnis"
    )
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    page_break(doc)

    # ----- 1. THE THREE SURFACES -----
    add_heading(doc, "1. There Are Three Claude Surfaces, Not One", level=1)
    para(
        doc,
        "Most people use the word \"Claude\" to mean one thing. There are actually three distinct products, "
        "and they have different architectures, different capabilities, and different right uses. Comparing them "
        "as if they were one product is the source of most of the confusion around what is or is not possible."
    )
    para(
        doc,
        "Claude.ai is the web product. You open it in a browser, you have a conversation, you can upload files "
        "into Projects, you can configure custom instructions per project. It has built-in retrieval augmented "
        "generation that kicks in when project knowledge approaches the context window limit, expanding effective "
        "capacity by roughly ten times. Memory persists per project. Cross-project memory does not exist by design."
    )
    para(
        doc,
        "Claude Desktop is the Mac and Windows desktop application. It is a hybrid surface. It can connect to "
        "local MCP servers, which means it can read and write files on your local machine through tool calls. "
        "It is closer to Claude Code than to Claude.ai, but it lacks the developer-grade lifecycle hooks, "
        "subprocess execution, and multi-instance coordination that Claude Code supports."
    )
    para(
        doc,
        "Claude Code is the command-line interface. It runs in a terminal in a working directory. It can read "
        "any file, write any file, execute shell commands, run hooks at lifecycle events, spawn sub-agents with "
        "their own models and tool allowlists, integrate with git, register arbitrary MCP servers, and coordinate "
        "work across multiple repositories. It is the most powerful surface, and it is the substrate that kipi "
        "is built on."
    )
    para(
        doc,
        "When this document says \"Claude.ai\" it means specifically the web product. When it says "
        "\"Claude Code\" it means specifically the CLI."
    )

    page_break(doc)

    # ----- 2. HOW CLAUDE.AI WORKS -----
    add_heading(doc, "2. How Claude.ai Works", level=1)
    para(
        doc,
        "Claude.ai is a conversational web product. The interaction model is straightforward: you type into a "
        "chat box, Claude responds in the same chat. Conversations can be organized into projects. Projects can "
        "have custom instructions and a knowledge base of uploaded files. Memory persists within a project across "
        "conversations."
    )
    para(
        doc,
        "Persistence in Claude.ai works through three mechanisms. First, project custom instructions: a block of "
        "text that gets prepended to every conversation in that project. Second, project knowledge: files you "
        "upload, which get indexed by Anthropic's servers and made retrievable. Third, conversation memory: "
        "claims and facts from past conversations within the project that Claude has been instructed to remember."
    )
    para(
        doc,
        "The RAG mechanism is the one that gets the most attention. When project knowledge exceeds the context "
        "window, the system automatically transitions from \"load everything into context\" to \"search and retrieve "
        "relevant chunks per query.\" This is invisible to the user; you simply have access to more knowledge "
        "than fits in a single context."
    )
    para(
        doc,
        "Claude.ai has also added MCP integration on the web side. Hosted MCP servers (like Notion, Gmail, "
        "Google Drive) can be authenticated and called from claude.ai conversations. This adds real capabilities "
        "to the web product: it can now read your inbox, query your databases, fetch from your Drive. The "
        "integrations are server-to-server through hosted MCP, not local subprocess execution."
    )
    para(
        doc,
        "What Claude.ai is genuinely good at: research, drafting, planning, analyzing pasted content, working "
        "with uploaded reference materials, conversations where the context fits inside one project, and "
        "interacting with hosted services through MCP integrations. The interaction is conversational, "
        "low-friction, and accessible from any browser without setup."
    )
    para(
        doc,
        "What Claude.ai cannot do: execute arbitrary shell commands on your local machine, edit files in your "
        "local filesystem directly (uploads are server-side; you cannot mutate your local repo through Claude.ai), "
        "run Python or Node scripts at lifecycle events, coordinate state across multiple isolated working "
        "directories, integrate with your local git, run subprocess hooks before and after tool calls, or load "
        "context conditionally from a markdown file tree on disk."
    )
    para(
        doc,
        "These limits are not bugs. They are intentional. The web product is shaped for safety, accessibility, "
        "and a chat-first user experience. Different shape, different job."
    )

    page_break(doc)

    # ----- 3. HOW CLAUDE CODE WORKS -----
    add_heading(doc, "3. How Claude Code Works", level=1)
    para(
        doc,
        "Claude Code is a CLI installed on the local machine. You launch it inside a working directory and it "
        "gives you a conversation with the model, with access to a different set of capabilities than the web "
        "product."
    )
    para(
        doc,
        "The core capabilities, simplified: read any file in or below the working directory, write any file, "
        "execute shell commands (subject to permissions), call MCP tools, spawn sub-agents, run hooks at lifecycle "
        "events, and integrate with git. The model has full access to the working directory. The user can grant "
        "or deny tool calls; the permissions model is settings-driven and audit-friendly."
    )
    para(
        doc,
        "The memory model in Claude Code is file-based. At session start, Claude Code reads CLAUDE.md from the "
        "working directory if it exists, plus an auto-memory directory at ~/.claude/projects/<project>/memory/ "
        "where individual memories are stored as markdown files with frontmatter declaring their type. There is "
        "no built-in RAG. Anthropic publicly stated that they tried RAG in Claude Code early on and dropped it "
        "in favor of agentic search (grep, glob, find) because agentic search outperformed RAG by a significant "
        "margin for code work."
    )
    para(
        doc,
        "This is the architectural fork in the road. Claude.ai built on RAG. Claude Code built on filesystem "
        "primitives. Both choices are defensible; they target different use cases."
    )
    para(
        doc,
        "Hooks are the most important capability in Claude Code that has no equivalent in Claude.ai. They are "
        "subprocess scripts wired in via .claude/settings.json to fire at specific lifecycle events: when the "
        "session starts, when the user submits a prompt, before any tool call, after any tool call, after the "
        "conversation is compacted, and at session stop. Each hook can read structured input from stdin, run "
        "arbitrary logic, write output that gets reinjected into the context, and exit with a code that affects "
        "the conversation flow. Hooks are how Claude Code becomes scriptable beyond the model's reasoning."
    )
    para(
        doc,
        "Sub-agents are the second capability without a Claude.ai equivalent. The orchestrator can spawn a "
        "sub-conversation with a different model, a different tool allowlist, a different system prompt, and "
        "a different effort level. The sub-agent runs to completion and returns a result. This allows different "
        "parts of a workflow to run on different models for cost and quality trade-offs, with each model seeing "
        "only the context relevant to its specific task."
    )
    para(
        doc,
        "Multi-instance work in Claude Code is just multiple working directories. Each directory is its own "
        "isolated workspace with its own CLAUDE.md, its own .claude/rules, its own canonical knowledge. Sessions "
        "opened in different directories cannot accidentally read each other's files unless the user explicitly "
        "grants cross-directory access. This is the foundation for running multiple businesses without state "
        "contamination."
    )

    page_break(doc)

    # ----- 4. ARCHITECTURE COMPARISON -----
    add_heading(doc, "4. Side-by-Side Architecture Comparison", level=1)
    para(
        doc,
        "The clearest way to see the difference is to look at the capabilities each surface provides and where "
        "they overlap or diverge."
    )
    add_table_rows(
        doc,
        ["Capability", "Claude.ai (web)", "Claude Code (CLI)"],
        [
            ["Conversational interface", "yes (browser)", "yes (terminal)"],
            ["File upload + RAG", "yes (auto-enables at context limit)", "no built-in RAG (agentic search via grep/glob)"],
            ["Local filesystem read/write", "no", "yes (any path in or below cwd)"],
            ["Custom instructions", "yes (per project)", "yes (CLAUDE.md, rules, output styles)"],
            ["Memory across sessions", "yes (per project, server-side)", "yes (CLAUDE.md + auto-memory, file-based)"],
            ["Cross-project memory", "no by design", "yes (via shared files or explicit cross-dir)"],
            ["MCP integrations", "yes (hosted MCP servers)", "yes (any MCP server including local subprocess)"],
            ["Local subprocess execution", "no", "yes (Bash, Python, any shell command)"],
            ["Lifecycle hooks", "no", "yes (SessionStart, PreToolUse, PostToolUse, Stop, etc.)"],
            ["Spawn sub-agents with different models", "no", "yes"],
            ["Git integration", "no (server-side)", "yes (full local git)"],
            ["Multi-repo / multi-instance coordination", "no", "yes (multiple working directories)"],
            ["Settings.json for permissions and behavior", "no", "yes"],
            ["Output styles", "no", "yes (.claude/output-styles/)"],
            ["Skills with auto-load triggers", "limited (via custom instructions)", "yes (rule-based auto-attach)"],
            ["Local privacy (files don't leave machine)", "no (uploads to Anthropic)", "yes (files stay on disk)"],
        ],
    )
    para(
        doc,
        "The pattern: Claude.ai is conversation-first with managed retrieval. Claude Code is filesystem-first "
        "with full local execution. The overlap is meaningful (both can use MCP, both can have custom "
        "instructions, both have memory), but the depth of integration is fundamentally different. Claude.ai "
        "is shaped for accessibility. Claude Code is shaped for control."
    )

    page_break(doc)

    # ----- 5. MEMORY COMPARISON -----
    add_heading(doc, "5. How Memory Works in Each Surface", level=1)
    para(
        doc,
        "The most common confusion in this comparison is around memory. Both surfaces have memory features; "
        "the architectures behind them are completely different."
    )
    para(
        doc,
        "Claude.ai memory is server-side and retrieval-based. When you upload a file to a project, it gets "
        "indexed by Anthropic's servers. When you have a conversation in that project, the system semantically "
        "searches the indexed content and retrieves the relevant chunks to include in the model's context. "
        "Above the context window limit, this is the only way you can have access to all the knowledge: by "
        "selectively retrieving the parts that match. Below the limit, the system can just load everything. "
        "The retrieval is probabilistic. The same question, asked with slightly different wording, can surface "
        "different chunks. This is good for fuzzy recall and bad for guaranteed access to specific information."
    )
    para(
        doc,
        "Conversation memory in Claude.ai works similarly but more aggressively. Claude can be instructed to "
        "remember facts from past conversations. The mechanism behind the scenes is also retrieval-based: the "
        "facts are stored, and Claude retrieves them when relevant queries come up. The user does not control "
        "directly what gets remembered or when."
    )
    para(
        doc,
        "Claude Code memory is file-based and deterministic. CLAUDE.md is loaded at session start, always, "
        "completely. Auto-memory is a directory of named markdown files with an index in MEMORY.md. The index "
        "is always loaded; individual memory files are pulled on demand when the index indicates they are "
        "relevant. .claude/rules/*.md files are conditionally attached based on path globs or content patterns "
        "declared in their frontmatter. Skills are loaded when explicitly invoked or when an auto-detection "
        "rule fires for them."
    )
    para(
        doc,
        "The difference in practice: if you ask Claude.ai a question, you do not know in advance what chunks "
        "of your project knowledge it will surface. The answer depends on what the embeddings match. If you "
        "ask Claude Code a question in a kipi instance, you know exactly which files were loaded, because the "
        "load rules are declared in frontmatter and the rule files themselves are on disk. The decision about "
        "what context the model sees is auditable and reproducible."
    )
    para(
        doc,
        "There is a real trade-off here. Claude.ai's RAG gives you access to large bodies of knowledge without "
        "you having to maintain any file structure. You just upload and ask. Claude Code's file-based memory "
        "requires that you (or your system) maintain enough discipline to give memories names, descriptions, "
        "and attach conditions that can be reasoned about. The discipline cost is real. What you get for it is "
        "predictability."
    )
    para(
        doc,
        "Both can be useful. They are not interchangeable."
    )

    page_break(doc)

    # ----- 6. WHAT KIPI ADDS -----
    add_heading(doc, "6. What Kipi Adds on Top of Claude Code", level=1)
    para(
        doc,
        "Kipi is not Claude Code. It is the operating system layer built on top of Claude Code. The previous "
        "sections compared the runtime to Claude.ai. This section explains why a stock Claude Code installation "
        "is also not sufficient for running multiple businesses, and what specifically kipi adds."
    )
    para(
        doc,
        "Out of the box, Claude Code gives you the capabilities listed above: filesystem access, hooks, "
        "sub-agents, MCP, git, multi-instance. But it does not give you a coherent way to use them. There are "
        "no enforced voice rules. There is no canonical knowledge structure. There is no bridge protocol "
        "between repositories. There is no gated workflow for product changes. There is no morning routine. "
        "There is no deterministic scorer for leads. The substrate is there; the operating system is not."
    )
    para(
        doc,
        "Kipi is the operating system. It consists of six plugins, roughly eighty-five Python tools registered "
        "as MCP, nineteen auto-loading rules, ten or so skills, five sub-agent definitions, six hook event "
        "handlers, thirty-five JSON schemas for inter-agent communication, eleven canonical knowledge files, "
        "and a CLI for propagating the skeleton across multiple instances. None of these come with Claude Code. "
        "They are the specific opinionation that makes a stock CLI installation into a multi-business operating "
        "environment."
    )
    para(
        doc,
        "The most important parts, in order of how often they fire:"
    )
    para(
        doc,
        "The voice enforcement chain runs on every piece of external-facing content the founder generates. A "
        "voice-enforcement rule auto-loads when external content is being drafted. The founder-voice skill is "
        "pulled in with its voice DNA references and writing samples. The model drafts. The kipi_voice_lint "
        "MCP tool runs Python regex against the draft, looking for banned phrases, AI tells, sentence-length "
        "uniformity, and emdashes. If the lint fails, the model rewrites. The check is structural, not "
        "instructional."
    )
    para(
        doc,
        "The wiring-check hook runs on every Edit or Write to the filesystem. It verifies that every change is "
        "connected end-to-end across plugins, hooks, MCP tools, agents, bus files, canonical files, and rules. "
        "This is what prevents kipi from accumulating dead code or dangling references. It runs as Python at "
        "PostToolUse, regardless of what the model thinks about whether the change is wired up."
    )
    para(
        doc,
        "The auto-detection rules watch for patterns in user input. A pasted meeting transcript triggers the "
        "debrief workflow without a command. A pasted social post screenshot triggers the engage workflow. A "
        "request to comment on someone else's content triggers the social-reaction-gate that forces extraction "
        "of the poster's claims before drafting. These are not the model deciding to be helpful. They are rules "
        "on disk that auto-attach when their trigger conditions match."
    )
    para(
        doc,
        "The multi-instance bridge is what makes parallel businesses possible. The KTLYST cluster has five "
        "role-specific instances (strategy, product, website, lawyer, personal-brand) sharing state through "
        "single-writer JSON files at ~/.ktlyst/bridge/. Strategy writes the canonical digest. Product writes "
        "the product state. Lawyer writes legal flags. No instance reads its own positioning from anywhere "
        "other than canonical-digest.json. The protocol prevents the same factual question from getting "
        "different answers in different sessions."
    )
    para(
        doc,
        "The PRD operating system enforces discipline on significant changes. Anything that would change a "
        "product or system goes through a state machine: idea, draft, in-review, approved, split, archived. "
        "Codex (a separate AI agent) runs both a native review and an adversarial review at each gate, "
        "streaming findings to a JSONL file. Every finding must be triaged with one of four dispositions "
        "before the PRD can advance. Concurrent PRD and issue contexts are blocked. The model cannot start "
        "implementation while a PRD is still being drafted, because the scope hook rejects edits outside the "
        "PRD file itself."
    )
    para(
        doc,
        "The deterministic Python core is what makes kipi trustable for numeric work. Lead scores load weights "
        "from a registry and compute by arithmetic. A/B test math is just math. Voice linting is regex against "
        "banned word lists. Cold-email validation walks structural rules. Churn health scores are deterministic "
        "feature aggregation. None of these involve the model. The model writes, decides which tool to call, "
        "and synthesizes; the deterministic core decides what counts as valid."
    )

    page_break(doc)

    # ----- 7. SCENARIOS -----
    add_heading(doc, "7. Practical Scenarios", level=1)
    para(
        doc,
        "Concrete examples are the fastest way to internalize the difference. Here are five recurring scenarios "
        "and how each surface handles them."
    )

    add_heading(doc, "7.1 Scenario: Founder pastes a meeting transcript and wants a debrief.", level=2)
    para(
        doc,
        "Claude.ai: founder pastes the transcript into the project chat. Claude reads it, generates a summary, "
        "possibly identifies the person and company, asks clarifying questions. The output stays in the chat. "
        "If the founder wants it saved, they have to copy it somewhere. The summary may or may not be retrievable "
        "later, depending on whether RAG surfaces it for the next relevant query. Relationship state is not "
        "tracked beyond what fits in the project memory."
    )
    para(
        doc,
        "Claude Code + kipi: the auto-detection rule fires when the transcript shape is recognized in the "
        "prompt. The /q-debrief workflow runs without a command. The structured template extracts the person, "
        "role, company, and intent. The 12 strategic implications lenses are run. Outputs are routed to canonical "
        "files: market intelligence goes to market-intelligence.md, talk tracks to talk-tracks.md, decisions to "
        "decisions.md. The follow-up loop tracker opens loops for next-actions. Notion is queued for a write. "
        "All of this happens in one prompt cycle, without the founder typing a command, because the rule on "
        "disk knows what to do with a transcript."
    )

    add_heading(doc, "7.2 Scenario: Founder needs to score a lead.", level=2)
    para(
        doc,
        "Claude.ai: founder asks Claude to score the lead based on the attributes. Claude returns a number. "
        "The number is generated by the model from the attributes and whatever scoring heuristics the model "
        "imagined during the response. Asking the same question with slightly different wording can produce a "
        "different score. The math is unauditable."
    )
    para(
        doc,
        "Claude Code + kipi: founder asks the model to score the lead. The model calls kipi_score_lead with "
        "the attributes and signals as JSON. The Python harness loads the weights from the scoring registry, "
        "applies them deterministically, returns a score and a top-reasons array. The same inputs always produce "
        "the same output. The weights are reviewable. If the founder disagrees with the score, they can change "
        "the weights file and rerun. The math is auditable end-to-end."
    )

    add_heading(doc, "7.3 Scenario: Founder is drafting a LinkedIn post.", level=2)
    para(
        doc,
        "Claude.ai: founder gives Claude the topic and asks for a post. Claude generates a draft based on its "
        "training and any custom instructions the founder has set. The draft might be in the founder's voice or "
        "might not. There is no structural check. The founder either accepts, edits, or rewrites."
    )
    para(
        doc,
        "Claude Code + kipi: the voice-enforcement rule auto-loads because external content is being drafted. "
        "The founder-voice skill is pulled in with its 400-line voice DNA and 600-line writing samples. The "
        "model drafts. kipi_voice_lint runs the draft through Python regex against banned phrases, AI tells, "
        "uniform sentence lengths, and emdashes. If the lint fails, specific violations are reported back to "
        "the model, which rewrites. The post does not ship until the check passes."
    )

    add_heading(doc, "7.4 Scenario: Founder running both a consulting practice and a startup.", level=2)
    para(
        doc,
        "Claude.ai: the founder creates two projects, one per business. Custom instructions are set per project. "
        "Knowledge files are uploaded per project. Conversation memory is siloed per project. The challenge: "
        "when something legitimately relevant to both businesses comes up (a contact, a market signal, a strategic "
        "decision), the founder has to manually copy it across. There is no cross-project memory by design."
    )
    para(
        doc,
        "Claude Code + kipi: each business is its own working directory with its own .claude/rules, its own "
        "canonical knowledge, its own voice variants where appropriate. Cross-business coordination happens "
        "through the bridge protocol. If a market signal is relevant to both, it lives in market_signal.json "
        "(written by whichever instance discovers it) and is readable from all relevant instances. The founder "
        "does not copy across; the protocol routes."
    )

    add_heading(doc, "7.5 Scenario: Founder making a serious product change.", level=2)
    para(
        doc,
        "Claude.ai: founder describes the change to Claude. Claude proposes an approach. Founder approves. "
        "Claude generates the code (if relevant tools are connected) or proposes the changes for the founder "
        "to make manually. There is no built-in adversarial review, no scope enforcement, no receipt-based "
        "closeout. The discipline is whatever the founder chooses to apply."
    )
    para(
        doc,
        "Claude Code + kipi: founder runs /prd-start with the idea. The PRD state machine begins. While the PRD "
        "is being drafted, scope enforcement restricts edits to the PRD file itself; implementation cannot "
        "begin yet. The founder runs /prd-review when ready. Codex runs both a native and an adversarial pass, "
        "writing findings to JSONL. Each finding gets a disposition. Only after all findings are triaged can "
        "/prd-approve run. The PRD is then split into issue specs. Each issue has an allowed_files list. The "
        "scope hook rejects edits outside the list during implementation. Verification, review, and findings "
        "triage must all complete before /issue-closeout. The discipline is structural."
    )

    page_break(doc)

    # ----- 8. WHEN TO USE WHICH -----
    add_heading(doc, "8. When Each Surface Is the Right Tool", level=1)
    para(
        doc,
        "Both surfaces have legitimate uses. The honest answer to \"which one should I use\" depends on what "
        "the work actually requires."
    )

    add_heading(doc, "8.1 Use Claude.ai when:", level=2)
    para(
        doc,
        "The work is conversational and the context fits in one project. The knowledge to consult lives in "
        "documents you can upload. You want zero setup. You want to access the assistant from any browser. "
        "You are comfortable with the model retrieving probabilistically from your project knowledge. The work "
        "is bounded by what one conversation can produce. You want the model to handle web-based integrations "
        "(Gmail, Drive, Notion) through hosted MCP rather than configuring local tools yourself."
    )
    para(
        doc,
        "Examples: drafting a long-form article from research you have already gathered. Synthesizing a policy "
        "document from uploaded references. Answering questions about a body of knowledge that fits in a project. "
        "Conducting deep research on a topic where the model can search and cite. Working alongside someone "
        "who wants to look over your shoulder."
    )

    add_heading(doc, "8.2 Use Claude Code + kipi when:", level=2)
    para(
        doc,
        "The work spans multiple sessions, multiple projects, or multiple businesses. State has to persist "
        "across sessions in a way you can audit. The work involves writing or editing files in a repository "
        "with discipline. You need deterministic checks (voice, scoring, validation) running on output before "
        "it ships. You need lifecycle hooks running Python at session boundaries. You need to coordinate work "
        "across multiple isolated working directories. Privacy matters and you want files to stay local. You "
        "want git-native operations. You want gated workflows for product changes."
    )
    para(
        doc,
        "Examples: running multiple businesses with separate voices and stakeholders. Maintaining a canonical "
        "knowledge base across years and hundreds of sessions. Generating content that must survive AI-detector "
        "patterns reliably. Making product or system changes that need adversarial review before merge. "
        "Coordinating role-specific instances (strategy, product, legal) that share some state. Tracking "
        "follow-up loops across conversations and platforms."
    )

    add_heading(doc, "8.3 Use both when:", level=2)
    para(
        doc,
        "The surfaces are not in opposition. A common pattern: research and exploration in Claude.ai (where "
        "RAG against a large uploaded corpus is convenient), then synthesis and execution in Claude Code + kipi "
        "(where outputs go through voice lint, state lands in canonical files, and follow-up loops get tracked). "
        "Or: brainstorming and rough drafts in Claude.ai during travel from a browser; structured writing and "
        "shipping in Claude Code + kipi from the laptop."
    )
    para(
        doc,
        "The key is to know what each surface is for. Use Claude.ai for tasks that map naturally to its shape: "
        "conversational, project-scoped, retrieval-friendly. Use Claude Code + kipi for tasks that need the "
        "operating-system properties: enforcement, persistence, isolation, multi-instance, deterministic. "
        "Forcing the wrong tool to do the other tool's job is where frustration comes from."
    )

    page_break(doc)

    # ----- 9. THE HONEST VERDICT -----
    add_heading(doc, "9. The Honest Verdict", level=1)
    para(
        doc,
        "Claude.ai is excellent at being a conversational AI with retrieval over your documents. Anthropic "
        "has invested heavily in making the project experience clean, the RAG transparent, and the integrations "
        "smooth. For the work it is shaped for, it is among the best products available."
    )
    para(
        doc,
        "Claude Code is excellent at being a local development environment for AI agents. It has the lifecycle "
        "hooks, the subprocess execution, the file system access, the sub-agent orchestration, and the MCP "
        "ecosystem to be the substrate for arbitrary local AI systems. Kipi is one such system."
    )
    para(
        doc,
        "Claude Code + kipi is excellent at being an operating environment for running multiple businesses "
        "with enforced voice, deterministic checks, multi-instance coordination, gated product workflows, and "
        "auditable persistence. It is what Claude Code can become when treated as a substrate rather than a "
        "chatbot."
    )
    para(
        doc,
        "None of these three are competitors to each other. They are different products solving different "
        "problems, and a sophisticated user picks whichever fits the work at hand. The framing where one of "
        "them is \"better\" than the others is the wrong frame. The right question is which tool's shape "
        "matches the work."
    )
    para(
        doc,
        "For the founder running multiple businesses today, the operating-system properties of Claude Code + "
        "kipi are what makes the day possible. Voice that survives every output without manual rewriting. "
        "Memory that compounds rather than overflows. Isolation between business contexts that prevents "
        "leakage. Deterministic checks that prevent numeric drift. Gated workflows that prevent reckless "
        "changes. These properties are not available in Claude.ai because they require local execution, "
        "hooks, and multi-instance coordination that the web product is not built to provide."
    )
    para(
        doc,
        "For somebody with a different shape of work, Claude.ai may be the right answer. Both are legitimate. "
        "The point of this document is to make the trade-offs visible so the choice can be made with eyes open, "
        "instead of choosing based on which name people have heard more often."
    )

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size} bytes")


if __name__ == "__main__":
    build()
